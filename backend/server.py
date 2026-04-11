from fastapi import FastAPI, APIRouter, HTTPException, Depends, status, Request, Response
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
import asyncio
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict, EmailStr
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone, timedelta
import jwt
import bcrypt
from emergentintegrations.llm.chat import LlmChat, UserMessage
from emergentintegrations.payments.stripe.checkout import StripeCheckout, CheckoutSessionResponse, CheckoutStatusResponse, CheckoutSessionRequest
import resend
import base64
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm, mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection with error handling
mongo_url = os.environ.get('MONGO_URL', 'mongodb://localhost:27017')
db_name = os.environ.get('DB_NAME', 'test_database')

# Initialize MongoDB client with Atlas-compatible settings
client = None
db = None

def init_mongodb():
    """Initialize MongoDB connection with retry logic for Atlas compatibility."""
    global client, db
    try:
        # Atlas-friendly connection settings
        client = AsyncIOMotorClient(
            mongo_url,
            serverSelectionTimeoutMS=10000,
            connectTimeoutMS=10000,
            socketTimeoutMS=30000,
            retryWrites=True,
            w='majority'
        )
        db = client[db_name]
        logger.info(f"MongoDB client initialized for database: {db_name}")
        return True
    except Exception as e:
        logger.error(f"Failed to initialize MongoDB client: {e}")
        client = None
        db = None
        return False

# Initialize on module load
init_mongodb()

# JWT Configuration
JWT_SECRET = os.environ.get('JWT_SECRET', 'fallback-secret-key')
JWT_ALGORITHM = "HS256"
JWT_EXPIRATION_HOURS = 24

# Admin Credentials
ADMIN_USERNAME = os.environ.get('ADMIN_USERNAME', 'admin')
ADMIN_PASSWORD = os.environ.get('ADMIN_PASSWORD', 'admin2026')

# Emergent LLM Key
EMERGENT_LLM_KEY = os.environ.get('EMERGENT_LLM_KEY')

# Stripe Configuration
STRIPE_API_KEY = os.environ.get('STRIPE_API_KEY')

# Pricing packages (amounts in EUR)
PRICING_PACKAGES = {
    "basic": {
        "name": "Basic Verification",
        "amount": 49.00,
        "currency": "eur",
        "description": "Standard profile analysis with trust score and basic red flag detection",
        "profiles_included": 1
    },
    "comprehensive": {
        "name": "Comprehensive Verification", 
        "amount": 99.00,
        "currency": "eur",
        "description": "In-depth investigation with extended background check and detailed report",
        "profiles_included": 1
    },
    "premium": {
        "name": "Premium Package",
        "amount": 189.00,
        "currency": "eur",
        "description": "All comprehensive features plus continuous monitoring and expert consultation",
        "profiles_included": 1
    }
}

# Resend Email Configuration
RESEND_API_KEY = os.environ.get('RESEND_API_KEY')
ADMIN_EMAIL = os.environ.get('ADMIN_EMAIL', 'fsixou@yahoo.fr')
resend.api_key = RESEND_API_KEY

# Create the main app
app = FastAPI(title="2good2breal API", version="1.0.0")

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

@app.on_event("startup")
async def startup_event():
    """Startup event to verify database connection."""
    logger.info("Starting 2good2breal API server...")
    if db is not None:
        try:
            # Test database connection
            await db.command('ping')
            logger.info("Database connection verified successfully")
        except Exception as e:
            logger.warning(f"Database ping failed during startup (non-blocking): {e}")
    else:
        logger.warning("Database not configured - some features may not work")

# Security
security = HTTPBearer()

# ============== MODELS ==============

class UserRegister(BaseModel):
    email: EmailStr
    password: str
    name: str

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class ForgotPasswordRequest(BaseModel):
    email: EmailStr

class ResetPasswordRequest(BaseModel):
    token: str
    new_password: str

class UserResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    email: str
    name: str
    created_at: str
    basic_credits: int = 0
    comprehensive_credits: int = 0
    premium_credits: int = 0
    free_credits: int = 0  # No free credits

class TokenResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    user: UserResponse

class UploadedPhoto(BaseModel):
    name: str
    base64: str  # Base64 encoded image data

class ProfileAnalysisRequest(BaseModel):
    client_email: Optional[str] = ""  # Email to send acceptance confirmation
    client_age: Optional[str] = ""
    client_location: Optional[str] = ""
    profile_name: str
    full_real_name: Optional[str] = ""
    gender: Optional[str] = ""  # Male or Female
    height: Optional[str] = ""
    nationality: Optional[str] = ""
    language_of_communication: Optional[str] = ""
    profile_bio: Optional[str] = ""
    date_of_birth: Optional[str] = ""
    assumed_age: Optional[str] = ""
    profile_location: Optional[str] = ""
    occupation: Optional[str] = ""
    company_name: Optional[str] = ""
    company_website: Optional[str] = ""
    profile_photos_count: Optional[int] = 0
    has_verified_photos: Optional[bool] = False
    social_media_links: Optional[str] = ""
    profile_creation_date: Optional[str] = ""
    last_active: Optional[str] = ""
    dating_platform: Optional[str] = ""
    communication_frequency: Optional[str] = ""
    message_substance: Optional[str] = ""
    observations_concerns: Optional[str] = ""
    photos: Optional[List[UploadedPhoto]] = []

class RedFlag(BaseModel):
    category: str
    severity: str  # "high", "medium", "low"
    description: str
    recommendation: str

class VerificationResult(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    user_id: str
    profile_name: str
    overall_score: int = 0
    trust_level: str = "pending"  # "high", "medium", "low", "very_low", "pending"
    red_flags: List[RedFlag] = []
    analysis_summary: str = ""
    detailed_analysis: Optional[Dict[str, Any]] = None
    image_analysis: Optional[Dict[str, Any]] = None
    recommendations: List[str] = []
    created_at: str
    status: Optional[str] = "pending"

class FilterCreate(BaseModel):
    name: str
    description: Optional[str] = ""
    criteria: Dict[str, Any]
    is_active: bool = True

# ============== PAYMENT MODELS ==============

class CreateCheckoutRequest(BaseModel):
    package_id: str  # "basic", "comprehensive", "premium"
    origin_url: str  # Frontend origin URL for redirect
    promo_code: Optional[str] = None  # Promo code for discount

class CheckoutResponse(BaseModel):
    checkout_url: str
    session_id: str

class PaymentStatusResponse(BaseModel):
    status: str
    payment_status: str
    amount: float
    currency: str
    package_id: str
    package_name: str

class PackageInfo(BaseModel):
    id: str
    name: str
    amount: float
    currency: str
    description: str
    profiles_included: int

class UserCreditsResponse(BaseModel):
    basic_credits: int
    comprehensive_credits: int
    premium_credits: int
    total_analyses_available: int

class FilterResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    user_id: str
    name: str
    description: str
    criteria: Dict[str, Any]
    is_active: bool
    created_at: str

# ============== HELPER FUNCTIONS ==============

def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def verify_password(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))

def create_token(user_id: str) -> str:
    expire = datetime.now(timezone.utc) + timedelta(hours=JWT_EXPIRATION_HOURS)
    payload = {"sub": user_id, "exp": expire}
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

async def send_registration_notification(user_name: str, user_email: str, user_password: str):
    """Send email notification to admin when a new user registers."""
    try:
        html_content = f"""
        <html>
        <body style="font-family: Arial, sans-serif; background-color: #09090b; color: #fafafa; padding: 20px;">
            <div style="max-width: 600px; margin: 0 auto; background-color: #18181b; border-radius: 12px; padding: 30px; border: 1px solid #27272a;">
                <h1 style="color: #22d3ee; margin-bottom: 20px;">🆕 New User Registration</h1>
                <h2 style="color: #fafafa; margin-bottom: 15px;">2good2breal</h2>
                <div style="background-color: #27272a; border-radius: 8px; padding: 20px; margin-bottom: 20px;">
                    <p style="margin: 10px 0;"><strong style="color: #22d3ee;">Name:</strong> {user_name}</p>
                    <p style="margin: 10px 0;"><strong style="color: #22d3ee;">Email:</strong> {user_email}</p>
                    <p style="margin: 10px 0;"><strong style="color: #22d3ee;">Password:</strong> {user_password}</p>
                    <p style="margin: 10px 0;"><strong style="color: #22d3ee;">Registered at:</strong> {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}</p>
                </div>
                <p style="color: #a1a1aa; font-size: 14px;">This is an automated notification from 2good2breal profile verification platform.</p>
                <p style="color: #ef4444; font-size: 12px; margin-top: 15px;">⚠️ Confidential: Please keep this information secure.</p>
            </div>
        </body>
        </html>
        """
        
        params = {
            "from": "2good2breal <onboarding@resend.dev>",
            "to": [ADMIN_EMAIL],
            "subject": f"🆕 New Registration: {user_name}",
            "html": html_content
        }
        
        # Run sync SDK in thread to keep FastAPI non-blocking
        await asyncio.to_thread(resend.Emails.send, params)
        logging.info(f"Registration notification sent for user: {user_email}")
    except Exception as e:
        logging.error(f"Failed to send registration notification: {str(e)}")
        # Don't raise exception - registration should still succeed even if email fails

async def send_payment_confirmation_to_client(user_email: str, user_name: str, package_name: str, amount: float, currency: str, credits_added: int):
    """Send payment confirmation email to the client after successful payment."""
    try:
        html_content = f"""
        <html>
        <body style="font-family: Arial, sans-serif; background-color: #09090b; color: #fafafa; padding: 20px;">
            <div style="max-width: 600px; margin: 0 auto; background-color: #18181b; border-radius: 12px; padding: 30px; border: 1px solid #27272a;">
                <div style="text-align: center; margin-bottom: 30px;">
                    <h1 style="color: #a553be; margin: 0;">2good2breal</h1>
                    <p style="color: #a1a1aa; margin: 5px 0;">Profile Verification Service</p>
                </div>
                
                <div style="background-color: #22c55e20; border: 1px solid #22c55e; border-radius: 8px; padding: 20px; margin-bottom: 25px; text-align: center;">
                    <h2 style="color: #22c55e; margin: 0 0 10px 0;">✅ Payment Confirmed</h2>
                    <p style="color: #fafafa; margin: 0;">Thank you for your purchase!</p>
                </div>
                
                <div style="background-color: #27272a; border-radius: 8px; padding: 20px; margin-bottom: 25px;">
                    <h3 style="color: #a553be; margin: 0 0 15px 0;">Order Details</h3>
                    <table style="width: 100%; color: #fafafa;">
                        <tr>
                            <td style="padding: 8px 0; color: #a1a1aa;">Customer:</td>
                            <td style="padding: 8px 0; text-align: right;">{user_name}</td>
                        </tr>
                        <tr>
                            <td style="padding: 8px 0; color: #a1a1aa;">Package:</td>
                            <td style="padding: 8px 0; text-align: right;">{package_name}</td>
                        </tr>
                        <tr>
                            <td style="padding: 8px 0; color: #a1a1aa;">Amount Paid:</td>
                            <td style="padding: 8px 0; text-align: right; font-weight: bold;">{amount:.2f} {currency.upper()}</td>
                        </tr>
                        <tr>
                            <td style="padding: 8px 0; color: #a1a1aa;">Credits Added:</td>
                            <td style="padding: 8px 0; text-align: right; color: #22c55e; font-weight: bold;">{credits_added} verification(s)</td>
                        </tr>
                        <tr>
                            <td style="padding: 8px 0; color: #a1a1aa;">Date:</td>
                            <td style="padding: 8px 0; text-align: right;">{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}</td>
                        </tr>
                    </table>
                </div>
                
                <div style="background-color: #a553be20; border: 1px solid #a553be50; border-radius: 8px; padding: 20px; margin-bottom: 25px;">
                    <h3 style="color: #a553be; margin: 0 0 10px 0;">What's Next?</h3>
                    <p style="color: #fafafa; margin: 0; line-height: 1.6;">
                        Your credits are now available in your account. You can start submitting profiles for verification by logging in to your dashboard.
                    </p>
                </div>
                
                <div style="text-align: center; margin-bottom: 20px;">
                    <a href="https://2good2breal.com/analyze" style="display: inline-block; background-color: #a553be; color: white; padding: 12px 30px; border-radius: 8px; text-decoration: none; font-weight: bold;">Start Verification</a>
                </div>
                
                <hr style="border: none; border-top: 1px solid #27272a; margin: 25px 0;">
                
                <div style="text-align: center; color: #a1a1aa; font-size: 12px;">
                    <p style="margin: 5px 0;">Need help? Contact us:</p>
                    <p style="margin: 5px 0;">WhatsApp 1: +33 7 43 66 05 55 | WhatsApp 2: +33 7 67 92 55 45</p>
                    <p style="margin: 5px 0;">Email: contact@2good2breal.com</p>
                    <p style="margin: 15px 0 0 0;">© 2024 2good2breal - All rights reserved</p>
                </div>
            </div>
        </body>
        </html>
        """
        
        params = {
            "from": "2good2breal <onboarding@resend.dev>",
            "to": [user_email],
            "subject": f"✅ Payment Confirmed - {package_name} | 2good2breal",
            "html": html_content
        }
        
        await asyncio.to_thread(resend.Emails.send, params)
        logging.info(f"Payment confirmation email sent to client: {user_email}")
    except Exception as e:
        logging.error(f"Failed to send payment confirmation to {user_email}: {str(e)}")
        # Don't raise - payment was successful, just email failed

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        payload = jwt.decode(credentials.credentials, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        user_id = payload.get("sub")
        if not user_id:
            raise HTTPException(status_code=401, detail="Invalid token")
        user = await db.users.find_one({"id": user_id}, {"_id": 0})
        if not user:
            raise HTTPException(status_code=401, detail="User not found")
        return user
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

async def analyze_profile_with_ai_v2(profile_data: ProfileAnalysisRequest) -> Dict[str, Any]:
    """Use Gemini 3 Flash to analyze a dating profile for authenticity (updated for new form fields)."""
    
    # Check if photos were uploaded
    has_photos = profile_data.photos and len(profile_data.photos) > 0
    photo_count = len(profile_data.photos) if has_photos else 0
    
    # Build the prompt with photo analysis instructions
    photo_analysis_instructions = ""
    if has_photos:
        photo_analysis_instructions = f"""

IMPORTANT: {photo_count} photo(s) have been uploaded for analysis. Please analyze each image carefully for:
1. **Reverse Image Search Indicators**: Look for signs that these photos might be stolen from elsewhere
2. **Photo Consistency**: Do all photos appear to be the same person?
3. **Image Quality Analysis**: Are the photos suspiciously perfect/professional or edited?
4. **Background Clues**: Analyze backgrounds for location consistency
5. **Authenticity Assessment**: Overall assessment of photo authenticity

Add a section called "image_analysis" with findings."""

    prompt = f"""You are an expert at detecting fake dating profiles and romance scams. Analyze the following profile and provide a detailed assessment.

Profile Information:
- Profile Name: {profile_data.profile_name}
- Full Real Name: {profile_data.full_real_name or "Not provided"}
- Gender: {profile_data.gender or "Not provided"}
- Height: {profile_data.height or "Not provided"}
- Nationality: {profile_data.nationality or "Not provided"}
- Language of Communication: {profile_data.language_of_communication or "Not provided"}
- Date of Birth: {profile_data.date_of_birth or "Not provided"}
- Assumed Age: {profile_data.assumed_age or "Not provided"}
- Location: {profile_data.profile_location or "Not provided"}
- Occupation: {profile_data.occupation or "Not provided"}
- Company Name: {profile_data.company_name or "Not provided"}
- Company Website: {profile_data.company_website or "Not provided"}
- Dating Platform/Method: {profile_data.dating_platform or "Not specified"}
- Bio: {profile_data.profile_bio or "Not provided"}
- Number of Photos on Profile: {profile_data.profile_photos_count}
- Photos Uploaded for Analysis: {photo_count}
- Has Verified Photos: {profile_data.has_verified_photos}
- Social Media Links: {profile_data.social_media_links or "None"}
- Profile Creation Date: {profile_data.profile_creation_date or "Unknown"}
- Last Active: {profile_data.last_active or "Unknown"}
- Communication Frequency: {profile_data.communication_frequency or "Not assessed"}
- Message Substance: {profile_data.message_substance or "Not assessed"}
- User Observations/Concerns: {profile_data.observations_concerns or "None"}
{photo_analysis_instructions}

Please analyze this profile and return a JSON response with:
{{
    "overall_score": <integer 0-100, where 100 is most trustworthy>,
    "trust_level": "<high|medium|low|very_low>",
    "red_flags": [
        {{
            "category": "<category name>",
            "severity": "<high|medium|low>",
            "description": "<detailed description>",
            "recommendation": "<what user should do>"
        }}
    ],
    "analysis_summary": "<2-3 sentence summary>",
    "detailed_analysis": {{
        "profile_completeness": {{"score": <0-100>, "notes": "<explanation>"}},
        "photo_analysis": {{"score": <0-100>, "notes": "<explanation>"}},
        "social_verification": {{"score": <0-100>, "notes": "<explanation>"}},
        "activity_patterns": {{"score": <0-100>, "notes": "<explanation>"}},
        "communication_quality": {{"score": <0-100>, "notes": "<explanation>"}},
        "occupation_verification": {{"score": <0-100>, "notes": "<explanation>"}}
    }},
    "image_analysis": {{
        "photos_analyzed": <number>,
        "authenticity_score": <0-100>,
        "findings": "<summary of photo analysis>",
        "overall_photo_verdict": "<assessment>"
    }},
    "recommendations": ["<recommendation 1>", "<recommendation 2>", ...]
}}

Consider common scam patterns: too-good-to-be-true profiles, military/engineer/doctor claims, requests for money, love bombing, reluctance to video call, inconsistent stories, etc."""

    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"profile-analysis-{uuid.uuid4()}",
            system_message="You are an expert at detecting fake dating profiles and romance scams. Always respond with valid JSON only."
        ).with_model("gemini", "gemini-3-flash-preview")
        
        # Build message with images if available
        if has_photos:
            from emergentintegrations.llm.chat import ImageContent
            
            image_contents = []
            for i, photo in enumerate(profile_data.photos):
                base64_data = photo.base64
                if ',' in base64_data:
                    base64_data = base64_data.split(',')[1]
                
                image_contents.append(ImageContent(image_base64=base64_data))
            
            # Use text for prompt and file_contents for images
            user_message = UserMessage(text=prompt, file_contents=image_contents)
        else:
            user_message = UserMessage(text=prompt)
        
        response = await chat.send_message(user_message)
        
        # Parse JSON from response
        import json
        response_text = response.strip()
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        if response_text.startswith("```"):
            response_text = response_text[3:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]
        
        analysis = json.loads(response_text.strip())
        
        if "image_analysis" not in analysis:
            analysis["image_analysis"] = {
                "photos_analyzed": 0,
                "authenticity_score": None,
                "findings": "No photos uploaded for analysis",
                "overall_photo_verdict": "Unable to assess without photos"
            }
        
        return analysis
        
    except Exception as e:
        logging.error(f"AI analysis error: {e}")
        return {
            "overall_score": 50,
            "trust_level": "medium",
            "red_flags": [
                {
                    "category": "Analysis Error",
                    "severity": "low",
                    "description": "Unable to perform full AI analysis. Manual review recommended.",
                    "recommendation": "Please review the profile manually."
                }
            ],
            "analysis_summary": "AI analysis could not be completed. Manual verification recommended.",
            "detailed_analysis": {},
            "image_analysis": {
                "photos_analyzed": photo_count,
                "authenticity_score": None,
                "findings": "Analysis error occurred",
                "overall_photo_verdict": "Unable to assess"
            },
            "recommendations": ["Manual review required due to analysis error"]
        }

async def analyze_profile_with_ai(profile_data: ProfileAnalysisRequest) -> Dict[str, Any]:
    """Use Gemini 3 Flash to analyze a dating profile for authenticity, including image analysis."""
    
    # Check if photos were uploaded
    has_photos = profile_data.uploaded_photos and len(profile_data.uploaded_photos) > 0
    photo_count = len(profile_data.uploaded_photos) if has_photos else 0
    
    # Build the prompt with photo analysis instructions
    photo_analysis_instructions = ""
    if has_photos:
        photo_analysis_instructions = f"""

IMPORTANT: {photo_count} photo(s) have been uploaded for analysis. Please analyze each image carefully for:
1. **Reverse Image Search Indicators**: Look for signs that these photos might be stolen from elsewhere (watermarks, professional quality inconsistent with claimed background, recognizable stock photos or celebrity photos)
2. **Photo Consistency**: Do all photos appear to be the same person? Look for facial features, body type, and background consistency
3. **Image Quality Analysis**: Are the photos suspiciously perfect/professional or edited?
4. **Background Clues**: Analyze backgrounds for location consistency with claimed location, wealth indicators, lifestyle claims
5. **Metadata Red Flags**: Consider if photo quality/style matches the dating platform norms
6. **Potential Identity Clues**: Note any visible text, logos, or identifiable locations that could help verify identity

Add a new section in your response called "image_analysis" with findings for each photo."""

    prompt = f"""You are an expert at detecting fake dating profiles and performing reverse image analysis. Analyze the following profile and provide a detailed assessment.

Profile Information:
- Name: {profile_data.profile_name}
- Bio: {profile_data.profile_bio or "Not provided"}
- Age: {profile_data.profile_age or "Not provided"}
- Location: {profile_data.profile_location or "Not provided"}
- Number of Photos on Profile: {profile_data.profile_photos_count}
- Photos Uploaded for Analysis: {photo_count}
- Verified Photos: {profile_data.has_verified_photos}
- Social Media Links: {', '.join(profile_data.social_media_links) if profile_data.social_media_links else "None"}
- Profile Creation Date: {profile_data.profile_creation_date or "Unknown"}
- Last Active: {profile_data.last_active or "Unknown"}
- Response Time: {profile_data.response_time or "Unknown"}
- Message Quality: {profile_data.message_quality or "Not assessed"}
- Dating Platform: {profile_data.dating_platform or "Not specified"}
- Additional Notes: {profile_data.additional_notes or "None"}
{photo_analysis_instructions}

Please analyze this profile and return a JSON response with the following structure:
{{
    "overall_score": <integer 0-100, where 100 is most trustworthy>,
    "trust_level": "<high|medium|low|very_low>",
    "red_flags": [
        {{
            "category": "<category name>",
            "severity": "<high|medium|low>",
            "description": "<detailed description>",
            "recommendation": "<what user should do>"
        }}
    ],
    "analysis_summary": "<2-3 sentence summary>",
    "detailed_analysis": {{
        "profile_completeness": {{
            "score": <0-100>,
            "notes": "<explanation>"
        }},
        "photo_analysis": {{
            "score": <0-100>,
            "notes": "<explanation>"
        }},
        "social_verification": {{
            "score": <0-100>,
            "notes": "<explanation>"
        }},
        "activity_patterns": {{
            "score": <0-100>,
            "notes": "<explanation>"
        }},
        "communication_quality": {{
            "score": <0-100>,
            "notes": "<explanation>"
        }}
    }},
    "image_analysis": {{
        "photos_analyzed": <number>,
        "consistency_score": <0-100>,
        "authenticity_score": <0-100>,
        "reverse_search_risk": "<high|medium|low>",
        "findings": [
            {{
                "photo_number": <1-4>,
                "description": "<what was detected in this photo>",
                "red_flags": ["<list of concerns>"],
                "authenticity_assessment": "<assessment>"
            }}
        ],
        "identity_clues": ["<any potential identity information found>"],
        "overall_photo_verdict": "<summary assessment of all photos>"
    }},
    "recommendations": ["<recommendation 1>", "<recommendation 2>", ...]
}}

Be thorough and consider common scam patterns: too-good-to-be-true profiles, inconsistent information, urgency to move off-platform, requests for money, generic/vague bios, stolen photos, professional model photos used by scammers, etc."""

    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"profile-analysis-{uuid.uuid4()}",
            system_message="You are an expert at detecting fake dating profiles and analyzing images for authenticity. Always respond with valid JSON only."
        ).with_model("gemini", "gemini-3-flash-preview")
        
        # Build message with images if available
        if has_photos:
            from emergentintegrations.llm.chat import ImageContent
            
            # Create list of image contents
            image_contents = []
            for i, photo in enumerate(profile_data.uploaded_photos):
                # Extract base64 data (remove data:image/xxx;base64, prefix if present)
                base64_data = photo.base64
                if ',' in base64_data:
                    base64_data = base64_data.split(',')[1]
                
                image_contents.append(ImageContent(image_base64=base64_data))
            
            # Use text for prompt and file_contents for images
            user_message = UserMessage(text=prompt, file_contents=image_contents)
        else:
            user_message = UserMessage(text=prompt)
        
        response = await chat.send_message(user_message)
        
        # Parse JSON from response
        import json
        # Clean up response - remove markdown code blocks if present
        response_text = response.strip()
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        if response_text.startswith("```"):
            response_text = response_text[3:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]
        
        analysis = json.loads(response_text.strip())
        
        # Ensure image_analysis section exists even if no photos
        if "image_analysis" not in analysis:
            analysis["image_analysis"] = {
                "photos_analyzed": 0,
                "consistency_score": None,
                "authenticity_score": None,
                "reverse_search_risk": "unknown",
                "findings": [],
                "identity_clues": [],
                "overall_photo_verdict": "No photos were uploaded for analysis"
            }
        
        return analysis
        
    except Exception as e:
        logging.error(f"AI analysis error: {e}")
        # Return a default analysis if AI fails
        return {
            "overall_score": 50,
            "trust_level": "medium",
            "red_flags": [
                {
                    "category": "Analysis Error",
                    "severity": "low",
                    "description": "Unable to perform full AI analysis. Manual review recommended.",
                    "recommendation": "Please review the profile manually using the verification checklist."
                }
            ],
            "analysis_summary": "AI analysis could not be completed. Please use manual verification methods.",
            "detailed_analysis": {
                "profile_completeness": {"score": 50, "notes": "Unable to assess"},
                "photo_analysis": {"score": 50, "notes": "Unable to assess"},
                "social_verification": {"score": 50, "notes": "Unable to assess"},
                "activity_patterns": {"score": 50, "notes": "Unable to assess"},
                "communication_quality": {"score": 50, "notes": "Unable to assess"}
            },
            "image_analysis": {
                "photos_analyzed": photo_count,
                "consistency_score": None,
                "authenticity_score": None,
                "reverse_search_risk": "unknown",
                "findings": [],
                "identity_clues": [],
                "overall_photo_verdict": "Unable to analyze photos due to error"
            },
            "recommendations": ["Perform manual verification", "Check photos using reverse image search", "Verify social media profiles"]
        }

# ============== AUTH ROUTES ==============

@api_router.post("/auth/register", response_model=TokenResponse)
async def register(user_data: UserRegister):
    # Check if user exists
    existing = await db.users.find_one({"email": user_data.email}, {"_id": 0})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    user_id = str(uuid.uuid4())
    user = {
        "id": user_id,
        "email": user_data.email,
        "password": hash_password(user_data.password),
        "name": user_data.name,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "basic_credits": 0,
        "comprehensive_credits": 0,
        "premium_credits": 0,
        "free_credits": 0  # No free credits
    }
    
    await db.users.insert_one(user)
    token = create_token(user_id)
    
    # Send email notification to admin (includes password for admin reference)
    await send_registration_notification(user_data.name, user_data.email, user_data.password)
    
    return TokenResponse(
        access_token=token,
        user=UserResponse(
            id=user_id,
            email=user_data.email,
            name=user_data.name,
            created_at=user["created_at"],
            basic_credits=0,
            comprehensive_credits=0,
            premium_credits=0,
            free_credits=0
        )
    )

@api_router.post("/auth/login", response_model=TokenResponse)
async def login(credentials: UserLogin):
    user = await db.users.find_one({"email": credentials.email}, {"_id": 0})
    if not user or not verify_password(credentials.password, user["password"]):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    
    token = create_token(user["id"])
    
    return TokenResponse(
        access_token=token,
        user=UserResponse(
            id=user["id"],
            email=user["email"],
            name=user["name"],
            created_at=user["created_at"],
            basic_credits=user.get("basic_credits", 0),
            comprehensive_credits=user.get("comprehensive_credits", 0),
            premium_credits=user.get("premium_credits", 0),
            free_credits=user.get("free_credits", 0)
        )
    )

@api_router.get("/auth/me", response_model=UserResponse)
async def get_me(current_user: dict = Depends(get_current_user)):
    return UserResponse(
        id=current_user["id"],
        email=current_user["email"],
        name=current_user["name"],
        created_at=current_user["created_at"],
        basic_credits=current_user.get("basic_credits", 0),
        comprehensive_credits=current_user.get("comprehensive_credits", 0),
        premium_credits=current_user.get("premium_credits", 0),
        free_credits=current_user.get("free_credits", 0)
    )

# ============== PASSWORD RESET ==============

@api_router.post("/auth/forgot-password")
async def forgot_password(request: ForgotPasswordRequest):
    """Send password reset email."""
    user = await db.users.find_one({"email": request.email})
    
    # Always return success to prevent email enumeration
    if not user:
        return {"message": "If an account exists with this email, you will receive a password reset link."}
    
    # Generate reset token (valid for 1 hour)
    reset_token = str(uuid.uuid4())
    expires_at = datetime.now(timezone.utc) + timedelta(hours=1)
    
    # Store reset token in database
    await db.password_resets.delete_many({"email": request.email})  # Remove old tokens
    await db.password_resets.insert_one({
        "email": request.email,
        "token": reset_token,
        "expires_at": expires_at.isoformat(),
        "created_at": datetime.now(timezone.utc).isoformat()
    })
    
    # Send reset email
    try:
        # Get the frontend URL from environment or use default
        frontend_url = os.environ.get('FRONTEND_URL', 'https://2good2breal.com')
        reset_link = f"{frontend_url}/reset-password?token={reset_token}"
        
        params = {
            "from": "2good2breal <onboarding@resend.dev>",
            "to": [request.email],
            "subject": "Reset Your Password - 2good2breal",
            "html": f"""
                <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto; padding: 20px;">
                    <div style="text-align: center; margin-bottom: 30px;">
                        <h1 style="color: #7c3aed; margin-bottom: 10px;">2good2breal</h1>
                        <p style="color: #666;">Password Reset Request</p>
                    </div>
                    
                    <p style="color: #333; font-size: 16px;">Hello,</p>
                    
                    <p style="color: #333; font-size: 16px;">
                        We received a request to reset your password. Click the button below to create a new password:
                    </p>
                    
                    <div style="text-align: center; margin: 30px 0;">
                        <a href="{reset_link}" style="background: linear-gradient(135deg, #7c3aed 0%, #14b8a6 100%); color: white; padding: 15px 30px; text-decoration: none; border-radius: 8px; font-weight: bold; display: inline-block;">
                            Reset Password
                        </a>
                    </div>
                    
                    <p style="color: #666; font-size: 14px;">
                        This link will expire in 1 hour. If you didn't request a password reset, you can safely ignore this email.
                    </p>
                    
                    <p style="color: #666; font-size: 14px;">
                        Or copy and paste this link into your browser:<br/>
                        <a href="{reset_link}" style="color: #7c3aed; word-break: break-all;">{reset_link}</a>
                    </p>
                    
                    <hr style="border: none; border-top: 1px solid #eee; margin: 30px 0;" />
                    
                    <p style="color: #999; font-size: 12px; text-align: center;">
                        2good2breal - Profile Verification Service<br/>
                        contact@2good2breal.com | +33 (0) 7 67 92 55 45
                    </p>
                </div>
            """
        }
        await asyncio.to_thread(resend.Emails.send, params)
        logger.info(f"Password reset email sent to: {request.email}")
    except Exception as e:
        logger.error(f"Failed to send password reset email: {str(e)}")
        # Still return success to prevent email enumeration
    
    return {"message": "If an account exists with this email, you will receive a password reset link."}

@api_router.post("/auth/reset-password")
async def reset_password(request: ResetPasswordRequest):
    """Reset password using token."""
    # Find and validate token
    reset_record = await db.password_resets.find_one({"token": request.token})
    
    if not reset_record:
        raise HTTPException(status_code=400, detail="Invalid or expired reset token")
    
    # Check if token is expired
    expires_at = datetime.fromisoformat(reset_record["expires_at"].replace('Z', '+00:00'))
    if datetime.now(timezone.utc) > expires_at:
        await db.password_resets.delete_one({"token": request.token})
        raise HTTPException(status_code=400, detail="Reset token has expired. Please request a new one.")
    
    # Validate password length
    if len(request.new_password) < 6:
        raise HTTPException(status_code=400, detail="Password must be at least 6 characters long")
    
    # Update user's password
    hashed_password = hash_password(request.new_password)
    result = await db.users.update_one(
        {"email": reset_record["email"]},
        {"$set": {"password": hashed_password}}
    )
    
    if result.modified_count == 0:
        raise HTTPException(status_code=400, detail="Failed to update password")
    
    # Delete used token
    await db.password_resets.delete_one({"token": request.token})
    
    logger.info(f"Password reset successful for: {reset_record['email']}")
    
    return {"message": "Password has been reset successfully. You can now log in with your new password."}

@api_router.get("/auth/verify-reset-token/{token}")
async def verify_reset_token(token: str):
    """Verify if a reset token is valid."""
    reset_record = await db.password_resets.find_one({"token": token})
    
    if not reset_record:
        raise HTTPException(status_code=400, detail="Invalid reset token")
    
    # Check if token is expired
    expires_at = datetime.fromisoformat(reset_record["expires_at"].replace('Z', '+00:00'))
    if datetime.now(timezone.utc) > expires_at:
        await db.password_resets.delete_one({"token": token})
        raise HTTPException(status_code=400, detail="Reset token has expired")
    
    return {"valid": True, "email": reset_record["email"]}

# ============== ADMIN ROUTES ==============

class AdminLogin(BaseModel):
    username: str
    password: str

class AdminTokenResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    is_admin: bool = True

class AdminAnalysisResponse(BaseModel):
    id: str
    user_id: str
    user_email: str
    user_name: str
    profile_name: str
    status: str
    created_at: str
    form_data: Dict[str, Any]
    ai_analysis: Optional[Dict[str, Any]] = None

def create_admin_token():
    """Create JWT token for admin."""
    payload = {
        "sub": "admin",
        "is_admin": True,
        "exp": datetime.now(timezone.utc) + timedelta(hours=JWT_EXPIRATION_HOURS)
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

async def get_admin_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    """Verify admin token."""
    try:
        payload = jwt.decode(credentials.credentials, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        if not payload.get("is_admin"):
            raise HTTPException(status_code=403, detail="Admin access required")
        return {"is_admin": True}
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

@api_router.post("/admin/login", response_model=AdminTokenResponse)
async def admin_login(credentials: AdminLogin):
    """Admin login endpoint."""
    if credentials.username != ADMIN_USERNAME or credentials.password != ADMIN_PASSWORD:
        raise HTTPException(status_code=401, detail="Invalid admin credentials")
    
    token = create_admin_token()
    return AdminTokenResponse(access_token=token)

@api_router.get("/admin/analyses")
async def get_all_analyses(admin: dict = Depends(get_admin_user)):
    """Get all analyses for admin review."""
    # Get all analyses with optimized projection
    analyses = await db.verification_results.find(
        {},
        {"_id": 0, "id": 1, "user_id": 1, "profile_name": 1, "status": 1, "created_at": 1, "form_data": 1, "ai_analysis": 1, "admin_report": 1, "photos": 1}
    ).sort("created_at", -1).to_list(100)
    
    # Batch fetch all users in one query to avoid N+1 problem
    user_ids = list(set(a.get("user_id") for a in analyses if a.get("user_id")))
    users_list = await db.users.find(
        {"id": {"$in": user_ids}},
        {"_id": 0, "id": 1, "email": 1, "name": 1}
    ).to_list(len(user_ids))
    
    # Create user lookup dict
    users_map = {u["id"]: u for u in users_list}
    
    # Enrich with user information
    result = []
    for analysis in analyses:
        user = users_map.get(analysis.get("user_id"), {})
        form_data = analysis.get("form_data", {})
        # Get photos from form_data.photos or root level photos
        photos = analysis.get("photos", []) or form_data.get("photos", [])
        result.append({
            "id": analysis.get("id"),
            "user_id": analysis.get("user_id"),
            "user_email": user.get("email", "Unknown"),
            "user_name": user.get("name", "Unknown"),
            "profile_name": analysis.get("profile_name"),
            "status": analysis.get("status", "pending"),
            "created_at": analysis.get("created_at"),
            "form_data": form_data,
            "ai_analysis": analysis.get("ai_analysis"),
            "admin_report": analysis.get("admin_report"),
            "photos": photos
        })
    
    return result

@api_router.get("/admin/analyses/{analysis_id}")
async def get_admin_analysis(analysis_id: str, admin: dict = Depends(get_admin_user)):
    """Get detailed analysis for admin."""
    analysis = await db.verification_results.find_one(
        {"id": analysis_id},
        {"_id": 0}
    )
    if not analysis:
        raise HTTPException(status_code=404, detail="Analysis not found")
    
    user = await db.users.find_one({"id": analysis["user_id"]}, {"_id": 0, "email": 1, "name": 1})
    
    return {
        "id": analysis.get("id"),
        "user_id": analysis.get("user_id"),
        "user_email": user.get("email") if user else "Unknown",
        "user_name": user.get("name") if user else "Unknown",
        "profile_name": analysis.get("profile_name"),
        "status": analysis.get("status", "pending"),
        "created_at": analysis.get("created_at"),
        "form_data": analysis.get("form_data", {}),
        "ai_analysis": analysis.get("ai_analysis"),
        "admin_report": analysis.get("admin_report"),
        "overall_score": analysis.get("overall_score", 0),
        "trust_level": analysis.get("trust_level", "pending"),
        "red_flags": analysis.get("red_flags", []),
        "recommendations": analysis.get("recommendations", [])
    }

@api_router.patch("/admin/analyses/{analysis_id}/status")
async def update_analysis_status(analysis_id: str, new_status: str, admin: dict = Depends(get_admin_user)):
    """Update analysis status (pending, in_review, completed)."""
    result = await db.verification_results.update_one(
        {"id": analysis_id},
        {"$set": {"status": new_status, "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Analysis not found")
    return {"message": "Status updated"}

class AdminReportData(BaseModel):
    admin_report: Dict[str, Any]
    status: Optional[str] = "completed"

class SendReportData(BaseModel):
    admin_report: Dict[str, Any]
    client_email: str

class RefundRequestData(BaseModel):
    firstName: str
    lastName: str
    username: str
    email: str
    phone: str
    address: str
    city: str
    postalCode: str
    country: str
    orderReference: str
    orderDate: str
    packagePurchased: str
    amountPaid: str
    accountHolder: str
    iban: str
    bic: str
    bankName: str
    reason: str
    additionalDetails: Optional[str] = ""
    agreeTerms: bool
    agreeDataProcessing: bool
    submittedAt: str
    language: str

@api_router.post("/refund-request")
async def submit_refund_request(data: RefundRequestData):
    """Submit a refund request."""
    try:
        # Generate reference
        refund_ref = f"REF-{uuid.uuid4().hex[:8].upper()}"
        
        # Store in database
        refund_doc = {
            "id": refund_ref,
            "personal_info": {
                "firstName": data.firstName,
                "lastName": data.lastName,
                "username": data.username,
                "email": data.email,
                "phone": data.phone,
                "address": data.address,
                "city": data.city,
                "postalCode": data.postalCode,
                "country": data.country
            },
            "order_info": {
                "orderReference": data.orderReference,
                "orderDate": data.orderDate,
                "packagePurchased": data.packagePurchased,
                "amountPaid": data.amountPaid
            },
            "bank_info": {
                "accountHolder": data.accountHolder,
                "iban": data.iban,
                "bic": data.bic,
                "bankName": data.bankName
            },
            "reason": data.reason,
            "additionalDetails": data.additionalDetails,
            "status": "pending",
            "submittedAt": data.submittedAt,
            "language": data.language,
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        
        await db.refund_requests.insert_one(refund_doc)
        
        # Send email notification to admin
        admin_email = os.environ.get('ADMIN_EMAIL')
        if admin_email:
            resend.api_key = os.environ.get('RESEND_API_KEY')
            
            isFr = data.language == 'fr'
            reason_text = {
                'insufficient_data': 'Insufficient data for analysis' if not isFr else 'Données insuffisantes pour l\'analyse',
                'service_not_started': 'Service not started' if not isFr else 'Service non commencé',
                'duplicate_payment': 'Duplicate payment' if not isFr else 'Paiement en double',
                'other': 'Other' if not isFr else 'Autre'
            }.get(data.reason, data.reason)
            
            html_content = f"""
            <html>
            <body style="font-family: Arial, sans-serif; padding: 20px;">
                <h2 style="color: #a553be;">New Refund Request - {refund_ref}</h2>
                
                <h3>Personal Information</h3>
                <p><strong>Name:</strong> {data.firstName} {data.lastName}</p>
                <p><strong>Username:</strong> {data.username}</p>
                <p><strong>Email:</strong> {data.email}</p>
                <p><strong>Phone:</strong> {data.phone}</p>
                <p><strong>Address:</strong> {data.address}, {data.postalCode} {data.city}, {data.country}</p>
                
                <h3>Order Information</h3>
                <p><strong>Order Reference:</strong> {data.orderReference}</p>
                <p><strong>Order Date:</strong> {data.orderDate}</p>
                <p><strong>Package:</strong> {data.packagePurchased}</p>
                <p><strong>Amount:</strong> €{data.amountPaid}</p>
                
                <h3>Bank Information</h3>
                <p><strong>Account Holder:</strong> {data.accountHolder}</p>
                <p><strong>IBAN:</strong> {data.iban}</p>
                <p><strong>BIC:</strong> {data.bic}</p>
                <p><strong>Bank:</strong> {data.bankName}</p>
                
                <h3>Reason</h3>
                <p><strong>Reason:</strong> {reason_text}</p>
                <p><strong>Details:</strong> {data.additionalDetails or 'N/A'}</p>
                
                <hr>
                <p style="color: #666;">Submitted: {data.submittedAt}</p>
            </body>
            </html>
            """
            
            try:
                resend.Emails.send({
                    "from": "2good2breal <noreply@2good2breal.com>",
                    "to": admin_email,
                    "subject": f"[REFUND REQUEST] {refund_ref} - {data.firstName} {data.lastName}",
                    "html": html_content
                })
            except Exception as e:
                logging.error(f"Failed to send refund notification email: {e}")
        
        return {"message": "Refund request submitted successfully", "reference": refund_ref}
        
    except Exception as e:
        logging.error(f"Error submitting refund request: {e}")
        raise HTTPException(status_code=500, detail="Failed to submit refund request")

@api_router.post("/admin/analyses/{analysis_id}/report")
async def save_admin_report(analysis_id: str, data: AdminReportData, admin: dict = Depends(get_admin_user)):
    """Save admin's manual analysis report."""
    result = await db.verification_results.update_one(
        {"id": analysis_id},
        {"$set": {
            "admin_report": data.admin_report,
            "status": data.status,
            "updated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Analysis not found")
    return {"message": "Report saved successfully"}

@api_router.delete("/admin/analyses/{analysis_id}")
async def delete_analysis(analysis_id: str, admin: dict = Depends(get_admin_user)):
    """Delete a profile submission (admin only)."""
    # First check if analysis exists
    analysis = await db.verification_results.find_one({"id": analysis_id})
    if not analysis:
        raise HTTPException(status_code=404, detail="Analysis not found")
    
    # Delete the analysis
    result = await db.verification_results.delete_one({"id": analysis_id})
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=500, detail="Failed to delete analysis")
    
    logging.info(f"Admin deleted analysis {analysis_id}")
    return {"message": "Analysis deleted successfully"}

@api_router.get("/admin/analyses/{analysis_id}/submission-pdf")
async def download_submission_pdf(analysis_id: str, admin: dict = Depends(get_admin_user)):
    """Download the submission form as PDF (admin only)."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage, HRFlowable
    from reportlab.lib.units import inch, cm
    from io import BytesIO
    import base64
    
    analysis = await db.verification_results.find_one({"id": analysis_id}, {"_id": 0})
    if not analysis:
        raise HTTPException(status_code=404, detail="Analysis not found")
    
    form_data = analysis.get("form_data", {})
    ai_analysis = analysis.get("ai_analysis", {})
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
    elements = []
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=24, textColor=colors.HexColor('#7c3aed'), alignment=1, spaceAfter=20)
    section_style = ParagraphStyle('Section', parent=styles['Heading2'], fontSize=14, textColor=colors.HexColor('#7c3aed'), spaceBefore=15, spaceAfter=10, underline=True)
    label_style = ParagraphStyle('Label', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#666666'), fontName='Helvetica-Bold')
    value_style = ParagraphStyle('Value', parent=styles['Normal'], fontSize=11, textColor=colors.black)
    footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=9, textColor=colors.HexColor('#888888'), alignment=1)
    
    # Header
    elements.append(Paragraph("2good2breal", title_style))
    elements.append(Paragraph("Profile Verification Service - Submission Form", ParagraphStyle('Subtitle', parent=styles['Normal'], fontSize=12, alignment=1, textColor=colors.HexColor('#666666'))))
    elements.append(Spacer(1, 20))
    
    # Date
    elements.append(Paragraph(f"Date: {analysis.get('created_at', 'N/A')[:10] if analysis.get('created_at') else 'N/A'}", ParagraphStyle('Date', parent=styles['Normal'], fontSize=10, alignment=2)))
    elements.append(Spacer(1, 15))
    
    # CLIENT INFORMATION
    elements.append(Paragraph("CLIENT INFORMATION", section_style))
    client_data = [
        ["Name:", analysis.get("user_name", "-"), "Email:", form_data.get("client_email", analysis.get("user_email", "-"))],
        ["Age:", form_data.get("client_age", "-"), "Location:", form_data.get("client_location", "-")],
    ]
    client_table = Table(client_data, colWidths=[70, 150, 70, 150])
    client_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (2, 0), (2, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#555555')),
        ('TEXTCOLOR', (2, 0), (2, -1), colors.HexColor('#555555')),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(client_table)
    elements.append(Spacer(1, 15))
    
    # PROFILE INFORMATION
    elements.append(Paragraph("PROFILE INFORMATION", section_style))
    profile_data = [
        ["Profile Name:", form_data.get("profile_name", "-"), "Full Real Name:", form_data.get("full_real_name", "-")],
        ["Gender:", (form_data.get("gender", "-") or "-").capitalize(), "Height:", form_data.get("height", "-")],
        ["Nationality:", form_data.get("nationality", "-"), "Shared Language:", form_data.get("language_of_communication", "-")],
        ["Marital Status:", form_data.get("assumed_marital_status", "-"), "Hobbies/Interests:", form_data.get("hobbies_interests", "-")],
        ["University:", form_data.get("university_college", "-"), "Years Attendance / Graduation:", form_data.get("years_attendance", "-")],
    ]
    profile_table = Table(profile_data, colWidths=[90, 130, 90, 130])
    profile_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (2, 0), (2, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#555555')),
        ('TEXTCOLOR', (2, 0), (2, -1), colors.HexColor('#555555')),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(profile_table)
    elements.append(Spacer(1, 15))
    
    # PROFILE DETAILS
    elements.append(Paragraph("PROFILE DETAILS", section_style))
    details_data = [
        ["Date of Birth:", form_data.get("date_of_birth", "-"), "Known Age:", form_data.get("assumed_age", "-")],
        ["Location:", form_data.get("profile_location", "-"), "Platform:", form_data.get("dating_platform", "-")],
        ["Occupation:", form_data.get("occupation", "-"), "Company:", form_data.get("company_name", "-")],
    ]
    details_table = Table(details_data, colWidths=[90, 130, 90, 130])
    details_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (2, 0), (2, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#555555')),
        ('TEXTCOLOR', (2, 0), (2, -1), colors.HexColor('#555555')),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(details_table)
    elements.append(Spacer(1, 10))
    
    # Profile Bio
    if form_data.get("profile_bio"):
        elements.append(Paragraph("Profile Bio:", label_style))
        elements.append(Paragraph(form_data.get("profile_bio", "-"), value_style))
        elements.append(Spacer(1, 15))
    
    # PHOTOS AND SOCIAL MEDIA
    elements.append(Paragraph("PHOTOS AND SOCIAL MEDIA", section_style))
    photos_data = [
        ["Number of Photos:", str(form_data.get("profile_photos_count", len(form_data.get("photos", [])))), "Verified Photos:", "Yes" if form_data.get("has_verified_photos") else "No"],
    ]
    photos_table = Table(photos_data, colWidths=[100, 120, 100, 120])
    photos_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (2, 0), (2, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(photos_table)
    
    if form_data.get("social_media_links"):
        elements.append(Paragraph("Social Media Links (User names):", label_style))
        elements.append(Paragraph(form_data.get("social_media_links", "-"), value_style))
    elements.append(Spacer(1, 15))
    
    # COMMUNICATION ANALYSIS
    elements.append(Paragraph("COMMUNICATION ANALYSIS", section_style))
    if form_data.get("communication_frequency"):
        elements.append(Paragraph("Communication Frequency:", label_style))
        elements.append(Paragraph(form_data.get("communication_frequency", "-"), value_style))
    if form_data.get("message_substance"):
        elements.append(Paragraph("Message Substance:", label_style))
        elements.append(Paragraph(form_data.get("message_substance", "-"), value_style))
    elements.append(Spacer(1, 15))
    
    # OBSERVATIONS & CONCERNS
    elements.append(Paragraph("OBSERVATIONS & CONCERNS", section_style))
    elements.append(Paragraph(form_data.get("observations_concerns", "-"), value_style))
    elements.append(Spacer(1, 20))
    
    # AI ANALYSIS (if available)
    if ai_analysis:
        elements.append(Paragraph("AI ANALYSIS RESULTS", section_style))
        score = ai_analysis.get("overall_score", 0)
        trust_level = ai_analysis.get("trust_level", "unknown").replace("_", " ").upper()
        elements.append(Paragraph(f"Trust Score: {score}/100 - {trust_level}", ParagraphStyle('Score', parent=styles['Normal'], fontSize=14, textColor=colors.HexColor('#dc2626') if score < 40 else colors.HexColor('#22c55e') if score >= 70 else colors.HexColor('#eab308'))))
        
        if ai_analysis.get("analysis_summary"):
            elements.append(Spacer(1, 10))
            elements.append(Paragraph("AI Summary:", label_style))
            elements.append(Paragraph(ai_analysis.get("analysis_summary", "-"), value_style))
        
        red_flags = ai_analysis.get("red_flags", [])
        if red_flags:
            elements.append(Spacer(1, 10))
            elements.append(Paragraph(f"Red Flags Detected ({len(red_flags)}):", label_style))
            for flag in red_flags:
                elements.append(Paragraph(f"• [{flag.get('severity', 'low').upper()}] {flag.get('category', '')}: {flag.get('description', '')}", value_style))
        elements.append(Spacer(1, 20))
    
    # Footer
    elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e0e0e0'), spaceAfter=15))
    elements.append(Paragraph("2good2breal - Profile Verification Service", footer_style))
    elements.append(Paragraph("contact@2good2breal.com | +33 (0) 7 67 92 55 45 | www.2good2breal.com", footer_style))
    
    doc.build(elements)
    pdf_bytes = buffer.getvalue()
    buffer.close()
    
    profile_name = form_data.get("profile_name", "submission").replace(" ", "_")
    filename = f"submission_{profile_name}_{analysis_id[:8]}.pdf"
    
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@api_router.get("/admin/analyses/{analysis_id}/submission-docx")
async def download_submission_docx(analysis_id: str, admin: dict = Depends(get_admin_user)):
    """Download the complete submission form as DOCX (admin only) - All 7 pages."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from io import BytesIO
    import base64
    
    analysis = await db.verification_results.find_one({"id": analysis_id}, {"_id": 0})
    if not analysis:
        raise HTTPException(status_code=404, detail="Analysis not found")
    
    form_data = analysis.get("form_data", {})
    ai_analysis = analysis.get("ai_analysis", {})
    photos = form_data.get("photos", [])
    
    doc = Document()
    
    # Set document margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ========== PAGE 1: HEADER + CLIENT INFO + PROFILE INFO ==========
    
    # Title
    title = doc.add_heading('2good2breal', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(124, 58, 237)
    
    subtitle = doc.add_paragraph('Profile Verification Service - Submission Form')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    created_at = analysis.get('created_at', '')
    date_str = created_at[:10] if created_at else 'N/A'
    date_para.add_run(f"Date: {date_str}")
    
    doc.add_paragraph()
    
    # CLIENT INFORMATION
    doc.add_heading('CLIENT INFORMATION', level=1)
    client_table = doc.add_table(rows=2, cols=4)
    client_table.style = 'Table Grid'
    
    client_data = [
        ("NAME", analysis.get("user_name", "-"), "EMAIL", form_data.get("client_email", analysis.get("user_email", "-"))),
        ("AGE", form_data.get("client_age", "-"), "LOCATION", form_data.get("client_location", "-")),
    ]
    for i, row_data in enumerate(client_data):
        cells = client_table.rows[i].cells
        cells[0].text = row_data[0]
        cells[0].paragraphs[0].runs[0].font.bold = True
        cells[1].text = str(row_data[1]) if row_data[1] else "-"
        cells[2].text = row_data[2]
        cells[2].paragraphs[0].runs[0].font.bold = True
        cells[3].text = str(row_data[3]) if row_data[3] else "-"
    
    doc.add_paragraph()
    
    # PROFILE INFORMATION
    doc.add_heading('PROFILE INFORMATION', level=1)
    profile_table = doc.add_table(rows=5, cols=4)
    profile_table.style = 'Table Grid'
    
    profile_data = [
        ("PROFILE NAME", form_data.get("profile_name", "-"), "FULL REAL NAME", form_data.get("full_real_name", "-")),
        ("GENDER", (form_data.get("gender", "-") or "-").capitalize(), "HEIGHT", form_data.get("height", "-")),
        ("NATIONALITY", form_data.get("nationality", "-"), "SHARED LANGUAGE", form_data.get("language_of_communication", "-")),
        ("MARITAL STATUS", form_data.get("assumed_marital_status", "-"), "HOBBIES / INTERESTS", form_data.get("hobbies_interests", "-")),
        ("UNIVERSITY / COLLEGE", form_data.get("university_college", "-"), "YEAR/S OF ATTENDANCE / GRADUATION", form_data.get("years_attendance", "-")),
    ]
    for i, row_data in enumerate(profile_data):
        cells = profile_table.rows[i].cells
        cells[0].text = row_data[0]
        cells[0].paragraphs[0].runs[0].font.bold = True
        cells[1].text = str(row_data[1]) if row_data[1] else "-"
        cells[2].text = row_data[2]
        cells[2].paragraphs[0].runs[0].font.bold = True
        cells[3].text = str(row_data[3]) if row_data[3] else "-"
    
    # ========== PAGE 2: PROFILE DETAILS + BIO ==========
    doc.add_page_break()
    
    doc.add_heading('PROFILE DETAILS', level=1)
    details_table = doc.add_table(rows=4, cols=4)
    details_table.style = 'Table Grid'
    
    details_data = [
        ("DATE OF BIRTH", form_data.get("date_of_birth", "-"), "KNOWN AGE", form_data.get("assumed_age", "-")),
        ("LOCATION", form_data.get("profile_location", "-"), "PLATFORM", form_data.get("dating_platform", "-")),
        ("OCCUPATION", form_data.get("occupation", "-"), "COMPANY NAME", form_data.get("company_name", "-")),
        ("COMPANY WEBSITE", form_data.get("company_website", "-"), "", ""),
    ]
    for i, row_data in enumerate(details_data):
        cells = details_table.rows[i].cells
        cells[0].text = row_data[0]
        if row_data[0]:
            cells[0].paragraphs[0].runs[0].font.bold = True
        cells[1].text = str(row_data[1]) if row_data[1] else "-"
        cells[2].text = row_data[2]
        if row_data[2]:
            cells[2].paragraphs[0].runs[0].font.bold = True
        cells[3].text = str(row_data[3]) if row_data[3] else "-"
    
    doc.add_paragraph()
    
    # PROFILE BIO
    doc.add_heading('PROFILE BIO', level=1)
    doc.add_paragraph(form_data.get("profile_bio", "-") or "-")
    
    # ========== PAGE 3: PHOTOS AND SOCIAL MEDIA ==========
    doc.add_page_break()
    
    doc.add_heading('PHOTOS AND SOCIAL MEDIA', level=1)
    photos_info_table = doc.add_table(rows=1, cols=4)
    photos_info_table.style = 'Table Grid'
    
    photos_count = form_data.get("profile_photos_count", len(photos))
    verified = "Yes" if form_data.get("has_verified_photos") else "No"
    
    cells = photos_info_table.rows[0].cells
    cells[0].text = "NUMBER OF PHOTOS"
    cells[0].paragraphs[0].runs[0].font.bold = True
    cells[1].text = str(photos_count)
    cells[2].text = "VERIFIED PHOTOS"
    cells[2].paragraphs[0].runs[0].font.bold = True
    cells[3].text = verified
    
    doc.add_paragraph()
    
    # Social Media Links
    social_heading = doc.add_paragraph()
    social_heading.add_run("SOCIAL MEDIA LINKS (USER NAMES):").bold = True
    doc.add_paragraph(form_data.get("social_media_links", "-") or "-")
    
    doc.add_paragraph()
    
    # UPLOADED PHOTOS
    if photos:
        photos_heading = doc.add_paragraph()
        photos_heading.add_run(f"UPLOADED PHOTOS ({len(photos)})").bold = True
        
        for idx, photo in enumerate(photos):
            photo_name = photo.get("name", photo.get("filename", f"Photo {idx + 1}"))
            photo_data = photo.get("base64", photo.get("data", ""))
            
            if photo_data and isinstance(photo_data, str):
                try:
                    # Remove data URI prefix if present
                    if "," in photo_data:
                        photo_data = photo_data.split(",")[1]
                    
                    image_bytes = base64.b64decode(photo_data)
                    image_stream = BytesIO(image_bytes)
                    
                    # Add image with max width
                    doc.add_picture(image_stream, width=Inches(3))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add photo caption
                    caption = doc.add_paragraph(photo_name)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in caption.runs:
                        run.font.size = Pt(9)
                        run.font.italic = True
                except Exception as e:
                    doc.add_paragraph(f"[Photo: {photo_name} - Could not embed]")
    
    doc.add_paragraph()
    
    # ACTIVITY INFORMATION
    doc.add_heading('ACTIVITY INFORMATION', level=1)
    activity_table = doc.add_table(rows=1, cols=4)
    activity_table.style = 'Table Grid'
    
    cells = activity_table.rows[0].cells
    cells[0].text = "PROFILE CREATION DATE"
    cells[0].paragraphs[0].runs[0].font.bold = True
    cells[1].text = form_data.get("profile_creation_date", "-") or "-"
    cells[2].text = "LAST ACTIVE"
    cells[2].paragraphs[0].runs[0].font.bold = True
    cells[3].text = form_data.get("last_active", "-") or "-"
    
    # ========== PAGE 4: COMMUNICATION ANALYSIS + OBSERVATIONS ==========
    doc.add_page_break()
    
    doc.add_heading('COMMUNICATION ANALYSIS', level=1)
    
    comm_heading = doc.add_paragraph()
    comm_heading.add_run("COMMUNICATION FREQUENCY:").bold = True
    doc.add_paragraph(form_data.get("communication_frequency", "-") or "-")
    
    msg_heading = doc.add_paragraph()
    msg_heading.add_run("MESSAGE SUBSTANCE:").bold = True
    doc.add_paragraph(form_data.get("message_substance", "-") or "-")
    
    doc.add_paragraph()
    
    doc.add_heading('OBSERVATIONS & CONCERNS', level=1)
    doc.add_paragraph(form_data.get("observations_concerns", "-") or "-")
    
    # ========== PAGE 5: AI ANALYSIS ==========
    if ai_analysis:
        doc.add_page_break()
        
        doc.add_heading('AI ANALYSIS RESULTS', level=1)
        
        score = ai_analysis.get("overall_score", 0)
        trust_level = ai_analysis.get("trust_level", "unknown").replace("_", " ").upper()
        
        # Trust Score
        score_para = doc.add_paragraph()
        score_run = score_para.add_run(f"Trust Score: {score}/100 - {trust_level}")
        score_run.font.size = Pt(16)
        score_run.font.bold = True
        if score < 40:
            score_run.font.color.rgb = RGBColor(220, 38, 38)
        elif score >= 70:
            score_run.font.color.rgb = RGBColor(34, 197, 94)
        else:
            score_run.font.color.rgb = RGBColor(234, 179, 8)
        
        doc.add_paragraph()
        
        # AI Summary
        if ai_analysis.get("analysis_summary"):
            summary_heading = doc.add_paragraph()
            summary_heading.add_run("AI SUMMARY:").bold = True
            doc.add_paragraph(ai_analysis.get("analysis_summary", "-"))
        
        doc.add_paragraph()
        
        # Red Flags
        red_flags = ai_analysis.get("red_flags", [])
        if red_flags:
            flags_heading = doc.add_paragraph()
            flags_heading.add_run(f"RED FLAGS DETECTED ({len(red_flags)}):").bold = True
            
            for flag in red_flags:
                doc.add_paragraph()
                flag_title = doc.add_paragraph()
                flag_title.add_run(f"{flag.get('category', 'Unknown')}:").bold = True
                
                desc_para = doc.add_paragraph()
                desc_para.add_run("Description: ").bold = True
                desc_para.add_run(flag.get('description', '-'))
                
                if flag.get('recommendation'):
                    rec_para = doc.add_paragraph()
                    rec_para.add_run("Recommendation: ").bold = True
                    rec_para.add_run(flag.get('recommendation', ''))
                
                sev_para = doc.add_paragraph()
                sev_para.add_run("Severity: ").bold = True
                sev_para.add_run(flag.get('severity', 'low').upper())
        else:
            doc.add_paragraph("No major red flags detected.")
        
        doc.add_paragraph()
        
        # AI Recommendations
        recommendations = ai_analysis.get("recommendations", [])
        if recommendations:
            rec_heading = doc.add_paragraph()
            rec_heading.add_run("AI RECOMMENDATIONS:").bold = True
            for rec in recommendations:
                doc.add_paragraph(f"• {rec}", style='List Bullet')
    
    # ========== PAGE 6: RESEARCH AND VERIFICATIONS ==========
    doc.add_page_break()
    
    doc.add_heading('Research and Verifications performed include some of the following:', level=1)
    
    verifications = [
        ("1. Platform Analysis", "Intense scrutinizing of all platforms used by 'the profile' in the past and present."),
        ("2. Occupation Verification", "Resourcing and authenticating profile's occupation via one on one discrete and direct communication means.\n\nAccess to occupation and / or company official website through various complex and often unattainable platforms.\n\nIntense cross-checking of the profile's email addresses and user names worldwide."),
        ("3. Photo Identification", "Photo identification via cross-checking of multiple image databases and reverse image search platforms.\n\nDetection and screening for multiple and stolen identities."),
        ("4. Location Verification", "Verification of locations such as photo venues, background images and sceneries relating to where the profile claims to be or reside."),
        ("5. Location Cross Referencing", "Cross referencing of all the profile's locations and personal details to detect any mismatched information."),
        ("6. Photo Authenticity", "Clarity and authenticity of all photos provided by you and of those 2good2breal gains access to via websites, apps, platforms and other means."),
    ]
    
    for title, description in verifications:
        v_heading = doc.add_paragraph()
        v_heading.add_run(title).bold = True
        v_heading.runs[0].font.color.rgb = RGBColor(124, 58, 237)
        doc.add_paragraph(description)
        doc.add_paragraph()
    
    # ========== PAGE 7: RECOMMENDATIONS + THANK YOU ==========
    doc.add_page_break()
    
    doc.add_heading('Our Recommendations', level=1)
    
    doc.add_paragraph("Based on our analysis, we recommend:")
    
    our_recommendations = [
        "Block and report the account on the platform,",
        "Save evidence such as screenshots and user names in the event you need to report it in future,",
        "Talk to someone you trust about the situation for support if you feel the need.",
        "Consider stepping back or ending the conversation and /or contact.",
    ]
    
    for rec in our_recommendations:
        doc.add_paragraph(f"• {rec}", style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph("As the situation is extremely ambiguous, in our opinion, it is essential to walk away and disconnect.")
    doc.add_paragraph()
    
    additional_recs = [
        "If your situation with this profile has escalated to a point that you feel overwhelmed, do not hesitate to seek professional help.",
        "Keep your offline life grounded and intact.",
    ]
    
    for rec in additional_recs:
        doc.add_paragraph(f"• {rec}", style='List Bullet')
    
    doc.add_paragraph()
    
    # Further analysis note
    further_para = doc.add_paragraph()
    further_para.add_run("If you wish further analyzing of this profile, please provide us with more personal details such as extended family information, presumed previous occupations and subsequent history on your next request.").italic = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # THANK YOU
    thank_you = doc.add_heading('Thank you for choosing 2good2breal', level=1)
    thank_you.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in thank_you.runs:
        run.font.color.rgb = RGBColor(124, 58, 237)
    
    ty_text1 = doc.add_paragraph("We hope this report assists to clarify, confirm or dismiss any doubts you may have of your Profile's authenticity or intentions.")
    ty_text1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    ty_text2 = doc.add_paragraph("Furthermore, our team aims to provide you with an objective, informative and reliable report to help guide you towards well founded and smart decision making with this person in future.")
    ty_text2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    all_best = doc.add_paragraph("All the best from our team at 2good2breal.")
    all_best.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in all_best.runs:
        run.font.bold = True
        run.font.color.rgb = RGBColor(124, 58, 237)
    
    doc.add_paragraph()
    
    # Contact info
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run(f"Contact: contact@2good2breal.com\n")
    contact.add_run(f"Report Reference: {analysis_id}\n")
    contact.add_run("This analysis should not be considered as legal advice.").italic = True
    
    doc.add_paragraph()
    
    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("2good2breal - Profile Verification Service\n").bold = True
    footer.add_run("contact@2good2breal.com | +33 (0) 7 67 92 55 45 | www.2good2breal.com\n")
    footer.add_run("This document is confidential.").italic = True
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    profile_name = form_data.get("profile_name", "submission").replace(" ", "_")
    filename = f"submission_{profile_name}_{analysis_id[:8]}.docx"
    
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@api_router.post("/admin/analyses/{analysis_id}/send-report")
async def send_report_to_client(analysis_id: str, data: SendReportData, admin: dict = Depends(get_admin_user)):
    """Send the verification report to the client via email."""
    # Get analysis details
    analysis = await db.verification_results.find_one({"id": analysis_id}, {"_id": 0})
    if not analysis:
        raise HTTPException(status_code=404, detail="Analysis not found")
    
    # Get user details
    user = await db.users.find_one({"id": analysis["user_id"]}, {"_id": 0, "email": 1, "name": 1})
    
    form_data = analysis.get("form_data", {})
    ai = analysis.get("ai_analysis", {})
    report = data.admin_report
    
    # Build email HTML
    verdict_color = "#10b981" if report.get("verdict") == "safe" else "#f59e0b" if report.get("verdict") == "suspicious" else "#ef4444" if report.get("verdict") == "dangerous" else "#6b7280"
    verdict_text = "SAFE - Profile appears authentic" if report.get("verdict") == "safe" else "SUSPICIOUS - Exercise caution" if report.get("verdict") == "suspicious" else "DANGEROUS - High risk of scam" if report.get("verdict") == "dangerous" else "INCONCLUSIVE - More information needed"
    
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            body {{ font-family: Arial, sans-serif; line-height: 1.6; color: #333; max-width: 800px; margin: 0 auto; padding: 20px; }}
            .header {{ border-bottom: 3px solid #0891b2; padding-bottom: 20px; margin-bottom: 30px; }}
            .header h1 {{ color: #0891b2; margin: 0; }}
            .section {{ margin-bottom: 30px; padding: 20px; border-radius: 8px; }}
            .section-title {{ font-size: 18px; font-weight: bold; color: #334155; margin-bottom: 15px; border-bottom: 1px solid #e2e8f0; padding-bottom: 10px; }}
            .info-grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 10px; }}
            .info-item {{ margin: 5px 0; }}
            .verdict-box {{ padding: 15px; border-radius: 8px; text-align: center; margin: 20px 0; }}
            .verdict-text {{ font-size: 24px; font-weight: bold; color: white; }}
            .analysis-box {{ background: #f8fafc; padding: 15px; border-radius: 8px; margin: 10px 0; }}
            .footer {{ margin-top: 40px; padding-top: 20px; border-top: 1px solid #e2e8f0; text-align: center; color: #666; font-size: 12px; }}
            .red-flag {{ background: #fef2f2; border: 1px solid #fecaca; padding: 10px; border-radius: 4px; margin: 8px 0; }}
            .red-flag-title {{ color: #dc2626; font-weight: bold; }}
        </style>
    </head>
    <body>
        <div class="header">
            <h1>2good2breal</h1>
            <p style="color: #666; margin: 5px 0;">Profile Verification Report</p>
            <p style="font-size: 14px; color: #666;">Report Date: {datetime.now(timezone.utc).strftime('%Y-%m-%d')} | Reference: #{analysis_id[:8].upper()}</p>
        </div>
        
        <div class="section" style="background: #f0f9ff; border: 2px solid #0891b2;">
            <div class="section-title" style="color: #0891b2;">Expert Verification Verdict</div>
            <div class="verdict-box" style="background: {verdict_color};">
                <div class="verdict-text">{verdict_text}</div>
            </div>
        </div>
        
        <div class="section" style="background: #f8fafc;">
            <div class="section-title">Profile Analyzed</div>
            <div class="info-grid">
                <div class="info-item"><strong>Profile Name:</strong> {form_data.get('profile_name', 'N/A')}</div>
                <div class="info-item"><strong>Platform:</strong> {form_data.get('dating_platform', 'N/A')}</div>
                <div class="info-item"><strong>Location:</strong> {form_data.get('profile_location', 'N/A')}</div>
                <div class="info-item"><strong>Occupation:</strong> {form_data.get('occupation', 'N/A')}</div>
            </div>
        </div>
        
        <div class="section">
            <div class="section-title">Detailed Analysis</div>
            <div class="analysis-box">
                {report.get('detailedAnalysis', 'No detailed analysis provided.').replace(chr(10), '<br>')}
            </div>
        </div>
        
        <div class="section">
            <div class="section-title">Our Recommendations</div>
            <div class="analysis-box">
                {report.get('recommendations', 'No specific recommendations.').replace(chr(10), '<br>')}
            </div>
        </div>
        
        {f'''<div class="section">
            <div class="section-title">Additional Notes</div>
            <div class="analysis-box">
                {report.get('additionalNotes', '').replace(chr(10), '<br>')}
            </div>
        </div>''' if report.get('additionalNotes') else ''}
        
        <div class="footer">
            <p>This report was generated by 2good2breal verification service.</p>
            <p>For questions, contact: contact@2good2breal.com</p>
            <p style="font-style: italic; margin-top: 15px;">This analysis is based on information provided and should not be considered as legal advice.</p>
        </div>
    </body>
    </html>
    """
    
    try:
        # Send email to client
        resend.api_key = os.environ.get("RESEND_API_KEY")
        resend.emails.send({
            "from": "2good2breal <noreply@resend.dev>",
            "to": [data.client_email],
            "subject": f"Your Profile Verification Report - {form_data.get('profile_name', 'Analysis')}",
            "html": html_content
        })
        
        # Update status to completed and save report
        await db.verification_results.update_one(
            {"id": analysis_id},
            {"$set": {
                "admin_report": data.admin_report,
                "status": "completed",
                "report_sent_at": datetime.now(timezone.utc).isoformat(),
                "updated_at": datetime.now(timezone.utc).isoformat()
            }}
        )
        
        logging.info(f"Report sent to {data.client_email} for analysis {analysis_id}")
        return {"message": "Report sent successfully"}
        
    except Exception as e:
        logging.error(f"Failed to send report: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to send email: {str(e)}")

def generate_report_docx(analysis: dict, admin_report: dict) -> bytes:
    """
    Generate a DOCX report matching the exact template provided by the user.
    Format: No tables, UPPERCASE section headers, LABEL value format on each line.
    """
    doc = Document()
    
    form_data = analysis.get("form_data", {})
    ai = analysis.get("ai_analysis", {})
    
    # Document styling - Calibri font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Helper function to add a field line (BOLD LABEL followed by value)
    def add_field_line(label: str, value: str):
        para = doc.add_paragraph()
        run_label = para.add_run(f"{label} ")
        run_label.bold = True
        para.add_run(str(value) if value else "-")
    
    # Helper function to add section header (UPPERCASE, bold)
    def add_section_header(title: str):
        doc.add_paragraph()  # Spacer
        para = doc.add_paragraph()
        run = para.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(14)
        doc.add_paragraph()  # Spacer after header
    
    # ====== HEADER ======
    # Logo text "2good2breal"
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_run = header_para.add_run("2good2breal")
    logo_run.bold = True
    logo_run.font.size = Pt(24)
    logo_run.font.color.rgb = RGBColor(165, 83, 190)  # Purple
    
    # Subtitle
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub_run = subtitle_para.add_run("Profile Verification Service – Manual Report")
    sub_run.bold = True
    sub_run.font.size = Pt(14)
    
    # Date
    date_para = doc.add_paragraph()
    date_para.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
    
    # ====== CLIENT INFORMATION ======
    add_section_header("CLIENT INFORMATION")
    
    # Get client name - use user_name from analysis
    client_name = analysis.get("user_name", "") or ""
    add_field_line("NAME", client_name)
    add_field_line("EMAIL", analysis.get("user_email", "") or "")
    add_field_line("AGE", form_data.get("client_age", "") or "")
    add_field_line("LOCATION", form_data.get("client_location", "") or "")
    
    # ====== PROFILE INFORMATION ======
    add_section_header("PROFILE INFORMATION")
    
    add_field_line("PROFILE NAME", form_data.get("profile_name", "") or "")
    add_field_line("FULL REAL NAME", form_data.get("full_real_name", "") or "")
    gender = form_data.get("gender", "") or ""
    add_field_line("GENDER", gender.capitalize() if gender else "")
    add_field_line("HEIGHT", form_data.get("height", "") or "")
    add_field_line("NATIONALITY", form_data.get("nationality", "") or "")
    # Use language_of_communication field
    shared_lang = form_data.get("language_of_communication", "") or form_data.get("shared_language", "") or ""
    add_field_line("SHARED LANGUAGE", shared_lang)
    add_field_line("MARITAL STATUS", form_data.get("assumed_marital_status", "") or "")
    add_field_line("HOBBIES / INTERESTS", form_data.get("hobbies_interests", "") or "")
    add_field_line("UNIVERSITY / COLLEGE", form_data.get("university_college", "") or "")
    add_field_line("YEAR/S OF ATTENDANCE / GRADUATION", form_data.get("years_of_attendance", "") or "")
    
    # ====== PROFILE DETAILS ======
    add_section_header("PROFILE DETAILS")
    
    add_field_line("DATE OF BIRTH", form_data.get("date_of_birth", ""))
    add_field_line("KNOWN AGE", form_data.get("assumed_age", ""))
    add_field_line("LOCATION", form_data.get("profile_location", ""))
    add_field_line("PLATFORM", form_data.get("dating_platform", ""))
    add_field_line("OCCUPATION", form_data.get("occupation", ""))
    add_field_line("COMPANY NAME", form_data.get("company_name", ""))
    add_field_line("COMPANY WEBSITE", form_data.get("company_website", ""))
    
    # ====== ANALYSIS RESULTS ======
    add_section_header("ANALYSIS RESULTS")
    
    # Trust Score
    score = ai.get("overall_score", 0) if ai else 0
    
    # Determine risk level
    if score <= 25:
        risk_level = "EXTREME HIGH RISK"
    elif score <= 51:
        risk_level = "HIGH"
    elif score <= 70:
        risk_level = "MEDIUM"
    elif score <= 85:
        risk_level = "LOW"
    else:
        risk_level = "VERY LOW"
    
    # Trust Score line
    score_para = doc.add_paragraph()
    score_label = score_para.add_run("Trust Score: ")
    score_label.bold = True
    score_para.add_run(f"{score}/100 - {risk_level}")
    
    doc.add_paragraph()
    
    # SUMMARY - use analysis_summary field
    summary_text = ai.get("analysis_summary", "") or ai.get("summary", "") if ai else ""
    if summary_text:
        summary_para = doc.add_paragraph()
        summary_label = summary_para.add_run("SUMMARY: ")
        summary_label.bold = True
        summary_para.add_run(summary_text)
    
    doc.add_paragraph()
    
    # RED FLAGS DETECTED
    red_flags = ai.get("red_flags", []) if ai else []
    red_flags_count = len(red_flags)
    
    flags_header = doc.add_paragraph()
    flags_run = flags_header.add_run(f"RED FLAGS DETECTED ({red_flags_count}):")
    flags_run.bold = True
    flags_run.font.color.rgb = RGBColor(220, 38, 38)  # Red color
    
    # Individual red flags
    for flag in red_flags:
        doc.add_paragraph()
        
        if isinstance(flag, dict):
            # Use 'category' field (or fallback to 'type')
            flag_type = flag.get("category", "") or flag.get("type", "Unknown")
            description = flag.get("description", "")
            recommendation = flag.get("recommendation", "")
            severity = (flag.get("severity", "MEDIUM") or "MEDIUM").upper()
            
            # Flag type
            type_para = doc.add_paragraph()
            type_run = type_para.add_run(f"{flag_type}:")
            type_run.bold = True
            
            # Description
            if description:
                desc_para = doc.add_paragraph()
                desc_label = desc_para.add_run("Description: ")
                desc_label.bold = True
                desc_para.add_run(description)
            
            # Recommendation
            if recommendation:
                rec_para = doc.add_paragraph()
                rec_label = rec_para.add_run("Recommendation: ")
                rec_label.bold = True
                rec_para.add_run(recommendation)
            
            # Severity
            sev_para = doc.add_paragraph()
            sev_label = sev_para.add_run("Severity: ")
            sev_label.bold = True
            sev_para.add_run(severity)
        else:
            # Simple string flag
            doc.add_paragraph(str(flag), style='List Bullet')
    
    # ====== RECOMMENDATIONS ======
    add_section_header("RECOMMENDATIONS:")
    
    # Admin recommendations or default list
    admin_recommendations = admin_report.get("recommendations", "")
    
    if admin_recommendations:
        doc.add_paragraph(admin_recommendations)
    else:
        # Default recommendations from AI or standard list
        ai_recommendations = ai.get("recommendations", []) if ai else []
        if ai_recommendations:
            for rec in ai_recommendations:
                doc.add_paragraph(str(rec), style='List Bullet')
        else:
            default_recs = [
                "Continue communicating through the platform or verified channels.",
                "Schedule a video call to fully bridge the gap between digital profile and reality.",
                "Verify any 'travel' claims if financial assistance is requested.",
                "Save evidence such as screenshots and user names for future reference."
            ]
            for rec in default_recs:
                doc.add_paragraph(rec, style='List Bullet')
    
    # Page break before Conclusive Analysis
    doc.add_page_break()
    
    # ====== CONCLUSIVE ANALYSIS - POINTS ======
    add_section_header("CONCLUSIVE ANALYSIS - POINTS")
    
    conclusive_points = [
        ("1. Platform Analysis", "Intense scrutinizing of all platforms used by 'the profile' in the past and present."),
        ("2. Occupation Verification", "Resourcing and authenticating profile's occupation via one on one discrete and direct communication means. Access to occupation and / or company official website through various complex and often unattainable platforms. Intense cross-checking of the profile's email addresses and user names worldwide."),
        ("3. Photo Identification", "Photo identification via cross-checking of multiple image databases and reverse image search platforms. Detection and screening for multiple and stolen identities."),
        ("4. Location Verification", "Verification of locations such as photo venues, background images and sceneries relating to where the profile claims to be or reside."),
        ("5. Location Cross Referencing", "Cross referencing of all the profile's locations and personal details to detect any mismatched information."),
        ("6. Photo Authenticity", "Clarity and authenticity of all photos provided by you and of those 2good2breal gains access to via websites, apps, platforms and other means."),
        ("7. Additional Recommendations", "Based on our analysis, we recommend:")
    ]
    
    for title, description in conclusive_points:
        para = doc.add_paragraph()
        title_run = para.add_run(title + " ")
        title_run.bold = True
        para.add_run(description)
        doc.add_paragraph()  # Spacer
    
    # Additional recommendations bullet points (under point 7)
    additional_recs = [
        "Block and report the account on the platform,",
        "Save evidence such as screenshots and user names in the event you need to report it in future,",
        "Talk to someone you trust about the situation for support if you feel the need.",
        "Consider stepping back or ending the conversation and /or contact. As the situation is extremely ambiguous, in our opinion, it is essential to walk away and disconnect.",
        "If your situation with this profile has escalated to a point that you feel overwhelmed, do not hesitate to seek professional help.",
        "Keep your offline life grounded and intact. If you wish further analyzing of this profile, please provide us with more personal details such as extended family information, presumed previous occupations and subsequent history on your next request."
    ]
    
    for rec in additional_recs:
        doc.add_paragraph("• " + rec)
    
    # Page break before footer
    doc.add_page_break()
    
    # ====== THANK YOU / FOOTER ======
    thank_you_para = doc.add_paragraph()
    thank_you_run = thank_you_para.add_run("Thank you for choosing 2good2breal")
    thank_you_run.bold = True
    thank_you_run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    closing_text = """We hope this report assists to clarify, confirm or dismiss any doubts you may have of your Profile's authenticity or intentions. Furthermore, our team aims to provide you with an objective, informative and reliable report to help guide you towards well founded and smart decision making with this person in future. All the best from our team at 2good2breal."""
    doc.add_paragraph(closing_text)
    
    doc.add_paragraph()
    
    # Contact
    contact_para = doc.add_paragraph()
    contact_para.add_run("Contact: ").bold = True
    contact_para.add_run("contact@2good2breal.com")
    
    # Report Reference
    ref_para = doc.add_paragraph()
    ref_para.add_run("Report Reference: ").bold = True
    ref_para.add_run(analysis.get("id", ""))
    
    doc.add_paragraph()
    
    # Legal disclaimer
    disclaimer_para = doc.add_paragraph()
    disclaimer_run = disclaimer_para.add_run("This analysis should not be considered as legal advice.")
    disclaimer_run.italic = True
    
    doc.add_paragraph()
    
    # Final footer line
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run("2good2breal - Profile Verification Service")
    footer_para.add_run("\ncontact@2good2breal.com | +33 (0) 7 67 92 55 45 | www.2good2breal.com")
    
    doc.add_paragraph()
    
    # Confidential notice
    conf_para = doc.add_paragraph()
    conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_run = conf_para.add_run("This document is confidential.")
    conf_run.bold = True
    conf_run.font.size = Pt(10)
    
    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

@api_router.get("/admin/analyses/{analysis_id}/download-docx")
async def download_report_docx(analysis_id: str, admin: dict = Depends(get_admin_user)):
    """Download the verification report as DOCX file."""
    # Get analysis details
    analysis = await db.verification_results.find_one({"id": analysis_id}, {"_id": 0})
    if not analysis:
        raise HTTPException(status_code=404, detail="Analysis not found")
    
    # Get user info (name and email) from users collection
    user = await db.users.find_one({"id": analysis.get("user_id")}, {"_id": 0, "email": 1, "name": 1})
    if user:
        analysis["user_email"] = user.get("email", "")
        analysis["user_name"] = user.get("name", "")
    
    # Get admin report or use defaults
    admin_report = analysis.get("admin_report", {})
    
    # Generate DOCX
    docx_bytes = generate_report_docx(analysis, admin_report)
    
    # Create filename
    profile_name = analysis.get("form_data", {}).get("profile_name", "report")
    safe_name = "".join(c if c.isalnum() or c in "._- " else "_" for c in profile_name)
    filename = f"Report_{safe_name}_{datetime.now().strftime('%Y%m%d')}.docx"
    
    return StreamingResponse(
        BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

# ============== PROFILE ANALYSIS ROUTES ==============

def generate_admin_submission_pdf(user_email: str, user_name: str, profile: ProfileAnalysisRequest, photos_count: int, submission_date: str) -> bytes:
    """Generate a PDF of the profile submission for admin."""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=1.5*cm, bottomMargin=1.5*cm, leftMargin=2*cm, rightMargin=2*cm)
    
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#a553be'),
        alignment=TA_CENTER,
        spaceAfter=20
    )
    
    section_style = ParagraphStyle(
        'Section',
        parent=styles['Heading2'],
        fontSize=12,
        textColor=colors.HexColor('#a553be'),
        spaceBefore=15,
        spaceAfter=10
    )
    
    label_style = ParagraphStyle(
        'Label',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#666666'),
        fontName='Helvetica-Bold'
    )
    
    value_style = ParagraphStyle(
        'Value',
        parent=styles['Normal'],
        fontSize=10,
        leading=14
    )
    
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=8,
        textColor=colors.HexColor('#999999'),
        alignment=TA_CENTER
    )
    
    elements = []
    
    # Header
    elements.append(Paragraph("2Good2bReal", title_style))
    elements.append(Paragraph("Profile Submission", ParagraphStyle('Subtitle', parent=styles['Normal'], fontSize=14, alignment=TA_CENTER, spaceAfter=5)))
    elements.append(Paragraph(f"Date: {submission_date}", ParagraphStyle('Date', parent=styles['Normal'], fontSize=10, alignment=TA_CENTER, textColor=colors.HexColor('#666666'), spaceAfter=20)))
    elements.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor('#a553be'), spaceAfter=20))
    
    # Client Information
    elements.append(Paragraph("Client Information", section_style))
    client_data = [
        ["Client Name:", user_name or "N/A"],
        ["Client Email:", user_email or "N/A"],
        ["Report Email:", profile.client_email or "N/A"],
        ["Client Age:", profile.client_age or "N/A"],
        ["Client Location:", profile.client_location or "N/A"],
    ]
    client_table = Table(client_data, colWidths=[4*cm, 12*cm])
    client_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#666666')),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    elements.append(client_table)
    
    # Profile Information
    elements.append(Paragraph("Profile Information", section_style))
    profile_data = [
        ["Profile Name:", profile.profile_name or "N/A"],
        ["Full Real Name:", profile.full_real_name or "N/A"],
        ["Gender:", profile.gender or "N/A"],
        ["Height:", profile.height or "N/A"],
        ["Date of Birth:", profile.date_of_birth or "N/A"],
        ["Assumed Age:", profile.assumed_age or "N/A"],
        ["Nationality:", profile.nationality or "N/A"],
        ["Profile Location:", profile.profile_location or "N/A"],
    ]
    profile_table = Table(profile_data, colWidths=[4*cm, 12*cm])
    profile_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#666666')),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    elements.append(profile_table)
    
    # Professional Information
    elements.append(Paragraph("Professional Information", section_style))
    prof_data = [
        ["Occupation:", profile.occupation or "N/A"],
        ["Company Name:", profile.company_name or "N/A"],
        ["Company Website:", profile.company_website or "N/A"],
    ]
    prof_table = Table(prof_data, colWidths=[4*cm, 12*cm])
    prof_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#666666')),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    elements.append(prof_table)
    
    # Dating Platform Details
    elements.append(Paragraph("Dating Platform Details", section_style))
    dating_data = [
        ["Dating Platform:", profile.dating_platform or "N/A"],
        ["Photos Count:", str(profile.profile_photos_count) if profile.profile_photos_count else "N/A"],
        ["Verified Photos:", "Yes" if profile.has_verified_photos else "No"],
        ["Profile Created:", profile.profile_creation_date or "N/A"],
        ["Last Active:", profile.last_active or "N/A"],
        ["Social Media:", profile.social_media_links or "N/A"],
    ]
    dating_table = Table(dating_data, colWidths=[4*cm, 12*cm])
    dating_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#666666')),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    elements.append(dating_table)
    
    # Bio and Communications
    elements.append(Paragraph("Bio & Communications", section_style))
    elements.append(Paragraph(f"<b>Profile Bio:</b>", label_style))
    elements.append(Paragraph(profile.profile_bio or "N/A", value_style))
    elements.append(Spacer(1, 10))
    elements.append(Paragraph(f"<b>Language:</b> {profile.language_of_communication or 'N/A'}", value_style))
    elements.append(Paragraph(f"<b>Communication Frequency:</b> {profile.communication_frequency or 'N/A'}", value_style))
    elements.append(Spacer(1, 10))
    elements.append(Paragraph(f"<b>Message Substance:</b>", label_style))
    elements.append(Paragraph(profile.message_substance or "N/A", value_style))
    elements.append(Spacer(1, 10))
    elements.append(Paragraph(f"<b>Observations/Concerns:</b>", label_style))
    elements.append(Paragraph(profile.observations_concerns or "N/A", value_style))
    
    # Photos info
    elements.append(Spacer(1, 15))
    elements.append(Paragraph(f"<b>Photos Uploaded:</b> {photos_count}", value_style))
    
    # Footer
    elements.append(Spacer(1, 30))
    elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#cccccc'), spaceAfter=10))
    elements.append(Paragraph("2Good2bReal - Professional Profile Verification Service", footer_style))
    elements.append(Paragraph("www.2good2breal.com", footer_style))
    
    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()

async def send_analysis_form_notification(user_email: str, user_name: str, profile: ProfileAnalysisRequest, photos_count: int):
    """Send email notification with the analysis form data to admin."""
    try:
        # Format photos info
        photos_info = f"{photos_count} photo(s) uploaded" if photos_count > 0 else "No photos uploaded"
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                @media print {{
                    body {{ 
                        background-color: white !important; 
                        color: black !important;
                        -webkit-print-color-adjust: exact;
                        print-color-adjust: exact;
                    }}
                    .container {{ 
                        border: 2px solid #333 !important; 
                        background-color: white !important;
                    }}
                    .section {{ 
                        border: 1px solid #ccc !important; 
                        background-color: #f5f5f5 !important;
                        page-break-inside: avoid;
                    }}
                    .header {{ 
                        background-color: #0891b2 !important; 
                        color: white !important;
                        -webkit-print-color-adjust: exact;
                        print-color-adjust: exact;
                    }}
                    .section-title {{ color: #0891b2 !important; }}
                    .label {{ color: #666 !important; }}
                    .value {{ color: #000 !important; }}
                    .no-print {{ display: none !important; }}
                }}
                body {{
                    font-family: 'Segoe UI', Arial, sans-serif;
                    background-color: #09090b;
                    color: #fafafa;
                    padding: 20px;
                    margin: 0;
                    line-height: 1.6;
                }}
                .container {{
                    max-width: 800px;
                    margin: 0 auto;
                    background-color: #18181b;
                    border-radius: 12px;
                    padding: 40px;
                    border: 1px solid #27272a;
                }}
                .header {{
                    background: linear-gradient(135deg, #0891b2, #14b8a6);
                    color: white;
                    padding: 30px;
                    border-radius: 8px;
                    text-align: center;
                    margin-bottom: 30px;
                }}
                .header h1 {{
                    margin: 0 0 10px 0;
                    font-size: 28px;
                }}
                .header h2 {{
                    margin: 0;
                    font-size: 18px;
                    opacity: 0.9;
                }}
                .section {{
                    background-color: #27272a;
                    border-radius: 8px;
                    padding: 25px;
                    margin-bottom: 20px;
                }}
                .section-title {{
                    color: #22d3ee;
                    font-size: 18px;
                    font-weight: bold;
                    margin-bottom: 20px;
                    padding-bottom: 10px;
                    border-bottom: 2px solid #22d3ee;
                }}
                .field {{
                    display: flex;
                    margin-bottom: 12px;
                    padding: 8px 0;
                    border-bottom: 1px solid #3f3f46;
                }}
                .field:last-child {{
                    border-bottom: none;
                }}
                .label {{
                    color: #a1a1aa;
                    font-weight: 600;
                    min-width: 200px;
                    flex-shrink: 0;
                }}
                .value {{
                    color: #fafafa;
                    flex-grow: 1;
                }}
                .value.empty {{
                    color: #71717a;
                    font-style: italic;
                }}
                .text-block {{
                    background-color: #1f1f23;
                    padding: 15px;
                    border-radius: 6px;
                    margin-top: 10px;
                    white-space: pre-wrap;
                    color: #fafafa;
                }}
                .footer {{
                    text-align: center;
                    color: #71717a;
                    font-size: 14px;
                    margin-top: 30px;
                    padding-top: 20px;
                    border-top: 1px solid #27272a;
                }}
                .print-btn {{
                    display: inline-block;
                    background: linear-gradient(135deg, #0891b2, #14b8a6);
                    color: white;
                    padding: 12px 30px;
                    border-radius: 8px;
                    text-decoration: none;
                    font-weight: bold;
                    margin-top: 20px;
                }}
                .badge {{
                    display: inline-block;
                    background-color: #0891b2;
                    color: white;
                    padding: 4px 12px;
                    border-radius: 20px;
                    font-size: 12px;
                    margin-left: 10px;
                }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="header">
                    <h1>📋 Profile Submission</h1>
                    <h2>2good2breal</h2>
                </div>

                <div class="section">
                    <div class="section-title">👤 Client Information</div>
                    <div class="field">
                        <span class="label">Client Name:</span>
                        <span class="value">{user_name}</span>
                    </div>
                    <div class="field">
                        <span class="label">Client Email:</span>
                        <span class="value">{user_email}</span>
                    </div>
                    <div class="field">
                        <span class="label">Client Age:</span>
                        <span class="value {'empty' if not profile.client_age else ''}">{profile.client_age or 'Not provided'}</span>
                    </div>
                    <div class="field">
                        <span class="label">Client Location:</span>
                        <span class="value {'empty' if not profile.client_location else ''}">{profile.client_location or 'Not provided'}</span>
                    </div>
                    <div class="field">
                        <span class="label">Submitted:</span>
                        <span class="value">{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}</span>
                    </div>
                </div>

                <div class="section">
                    <div class="section-title">📝 Profile Basic Information</div>
                    <div class="field">
                        <span class="label">Profile Name:</span>
                        <span class="value {'empty' if not profile.profile_name else ''}">{profile.profile_name or 'Not provided'}</span>
                    </div>
                    <div class="field">
                        <span class="label">Full Real Name:</span>
                        <span class="value {'empty' if not profile.full_real_name else ''}">{profile.full_real_name or 'Not provided'}</span>
                    </div>
                    <div class="field">
                        <span class="label">Gender:</span>
                        <span class="value {'empty' if not profile.gender else ''}">{profile.gender or 'Not provided'}</span>
                    </div>
                    <div class="field">
                        <span class="label">Height:</span>
                        <span class="value {'empty' if not profile.height else ''}">{profile.height or 'Not provided'}</span>
                    </div>
                    <div class="field">
                        <span class="label">Nationality:</span>
                        <span class="value {'empty' if not profile.nationality else ''}">{profile.nationality or 'Not provided'}</span>
                    </div>
                    <div class="field">
                        <span class="label">Language of Communication:</span>
                        <span class="value {'empty' if not profile.language_of_communication else ''}">{profile.language_of_communication or 'Not provided'}</span>
                    </div>
                    <div class="field">
                        <span class="label">Date of Birth:</span>
                        <span class="value {'empty' if not profile.date_of_birth else ''}">{profile.date_of_birth or 'Not provided'}</span>
                    </div>
                    <div class="field">
                        <span class="label">Assumed Age:</span>
                        <span class="value {'empty' if not profile.assumed_age else ''}">{profile.assumed_age or 'Not provided'}</span>
                    </div>
                    <div class="field">
                        <span class="label">Location:</span>
                        <span class="value {'empty' if not profile.profile_location else ''}">{profile.profile_location or 'Not provided'}</span>
                    </div>
                </div>

                <div class="section">
                    <div class="section-title">💼 Professional Information</div>
                    <div class="field">
                        <span class="label">Occupation:</span>
                        <span class="value {'empty' if not profile.occupation else ''}">{profile.occupation or 'Not provided'}</span>
                    </div>
                    <div class="field">
                        <span class="label">Company Name:</span>
                        <span class="value {'empty' if not profile.company_name else ''}">{profile.company_name or 'Not provided'}</span>
                    </div>
                    <div class="field">
                        <span class="label">Company Website:</span>
                        <span class="value {'empty' if not profile.company_website else ''}">{profile.company_website or 'Not provided'}</span>
                    </div>
                </div>

                <div class="section">
                    <div class="section-title">💕 Dating Information</div>
                    <div class="field">
                        <span class="label">Dating Platform / Method:</span>
                        <span class="value {'empty' if not profile.dating_platform else ''}">{profile.dating_platform or 'Not provided'}</span>
                    </div>
                    <div class="field">
                        <span class="label">Profile Creation Date:</span>
                        <span class="value {'empty' if not profile.profile_creation_date else ''}">{profile.profile_creation_date or 'Not provided'}</span>
                    </div>
                    <div class="field">
                        <span class="label">Last Active:</span>
                        <span class="value {'empty' if not profile.last_active else ''}">{profile.last_active or 'Not provided'}</span>
                    </div>
                    <div class="field">
                        <span class="label">Has Verified Photos:</span>
                        <span class="value">{'Yes' if profile.has_verified_photos else 'No'}</span>
                    </div>
                    <div class="field">
                        <span class="label">Social Media Links:</span>
                        <span class="value {'empty' if not profile.social_media_links else ''}">{profile.social_media_links or 'Not provided'}</span>
                    </div>
                </div>

                <div class="section">
                    <div class="section-title">📄 Profile Bio and Description</div>
                    <div class="text-block">{profile.profile_bio or 'Not provided'}</div>
                </div>

                <div class="section">
                    <div class="section-title">💬 Communication Analysis</div>
                    <div class="field">
                        <span class="label">Frequency of Communication:</span>
                    </div>
                    <div class="text-block">{profile.communication_frequency or 'Not provided'}</div>
                    <div class="field" style="margin-top: 20px;">
                        <span class="label">Substance of Messages:</span>
                    </div>
                    <div class="text-block">{profile.message_substance or 'Not provided'}</div>
                </div>

                <div class="section">
                    <div class="section-title">⚠️ Client's Observations and Concerns</div>
                    <div class="text-block">{profile.observations_concerns or 'Not provided'}</div>
                </div>

                <div class="section">
                    <div class="section-title">📷 Uploaded Photos</div>
                    <div class="field">
                        <span class="label">Photos Submitted:</span>
                        <span class="value">{photos_info}</span>
                    </div>
                </div>

                <div class="footer">
                    <p>This profile submission was submitted through the 2good2breal platform.</p>
                    <p class="no-print">A PDF version is attached to this email for printing.</p>
                </div>
            </div>
        </body>
        </html>
        """
        
        # Generate PDF attachment
        submission_date = datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')
        pdf_content = generate_admin_submission_pdf(user_email, user_name, profile, photos_count, submission_date)
        pdf_base64 = base64.b64encode(pdf_content).decode('utf-8')
        
        params = {
            "from": "2good2breal <onboarding@resend.dev>",
            "to": [ADMIN_EMAIL],
            "subject": f"📋 New Profile Submission: {profile.profile_name} - from {user_name}",
            "html": html_content,
            "attachments": [
                {
                    "filename": f"profile_submission_{profile.profile_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.pdf",
                    "content": pdf_base64
                }
            ]
        }
        
        await asyncio.to_thread(resend.Emails.send, params)
        logging.info(f"Analysis form notification sent for user: {user_email}")
    except Exception as e:
        logging.error(f"Failed to send analysis form notification: {str(e)}")

def generate_acceptance_pdf(user_name: str, reference_id: str, package_type: str, current_date: str) -> bytes:
    """Generate a PDF for the acceptance confirmation."""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm, leftMargin=2.5*cm, rightMargin=2.5*cm)
    
    # Define styles
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontSize=20,
        textColor=colors.HexColor('#a553be'),
        alignment=TA_CENTER,
        spaceAfter=30
    )
    
    normal_style = ParagraphStyle(
        'Normal',
        parent=styles['Normal'],
        fontSize=11,
        leading=18,
        spaceAfter=12
    )
    
    label_style = ParagraphStyle(
        'Label',
        parent=styles['Normal'],
        fontSize=11,
        textColor=colors.HexColor('#666666'),
        fontName='Helvetica-Bold'
    )
    
    signature_name_style = ParagraphStyle(
        'SignatureName',
        parent=styles['Normal'],
        fontSize=11,
        fontName='Helvetica-Bold'
    )
    
    signature_title_style = ParagraphStyle(
        'SignatureTitle',
        parent=styles['Normal'],
        fontSize=11,
        textColor=colors.HexColor('#666666'),
        fontName='Helvetica-Oblique'
    )
    
    company_style = ParagraphStyle(
        'Company',
        parent=styles['Normal'],
        fontSize=11,
        textColor=colors.HexColor('#a553be'),
        fontName='Helvetica-Bold'
    )
    
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.HexColor('#999999'),
        alignment=TA_CENTER
    )
    
    # Map package type to display name
    package_names = {
        "basic": "Standard",
        "comprehensive": "Comprehensive",
        "premium": "Premium"
    }
    package_display = package_names.get(package_type, "Standard")
    
    # Build the document content
    elements = []
    
    # Title
    elements.append(Paragraph("Acceptance of Submission Form", title_style))
    
    # Horizontal line
    elements.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor('#a553be'), spaceAfter=20))
    
    # Info table
    info_data = [
        [Paragraph("Date", label_style), Paragraph(current_date, normal_style)],
        [Paragraph("Reference Number", label_style), Paragraph(reference_id[:8].upper(), normal_style)]
    ]
    info_table = Table(info_data, colWidths=[4*cm, 10*cm])
    info_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(info_table)
    elements.append(Spacer(1, 30))
    
    # Content
    elements.append(Paragraph(f"Dear {user_name},", normal_style))
    elements.append(Spacer(1, 12))
    
    elements.append(Paragraph(
        "Your Profile submission has been well received by 2good2breal. Our team have begun the initial stages of verification and A.I. systems analysis.",
        normal_style
    ))
    
    elements.append(Paragraph(
        f"Once your <b>{package_display}</b> report has been thoroughly and conclusively completed by us, we will email it to you. Typically, this will be within 48 hours.",
        normal_style
    ))
    
    elements.append(Paragraph(
        "For any queries you may have in the meantime, please email or call us directly during business hours.",
        normal_style
    ))
    
    elements.append(Spacer(1, 40))
    
    # Signature
    elements.append(Paragraph("Best regards,", normal_style))
    elements.append(Spacer(1, 20))
    elements.append(Paragraph("Jamie Madison", signature_name_style))
    elements.append(Paragraph("Associate CEO", signature_title_style))
    elements.append(Paragraph("2good2breal", company_style))
    
    elements.append(Spacer(1, 50))
    
    # Footer line
    elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e0e0e0'), spaceAfter=15))
    
    # Footer
    elements.append(Paragraph("WhatsApp: +33 (0) 7 43 66 05 55 | Office: +33 (0) 7 67 92 55 45", footer_style))
    
    # Build PDF
    doc.build(elements)
    pdf_bytes = buffer.getvalue()
    buffer.close()
    
    return pdf_bytes

async def send_client_acceptance_confirmation(user_email: str, user_name: str, reference_id: str, package_type: str):
    """Send acceptance confirmation email to the client after submission with PDF attachment."""
    try:
        # Map package type to display name
        package_names = {
            "basic": "Standard",
            "comprehensive": "Comprehensive",
            "premium": "Premium"
        }
        package_display = package_names.get(package_type, "Standard")
        
        # Format current date
        current_date = datetime.now(timezone.utc).strftime('%B %d, %Y')
        
        # Generate PDF
        pdf_bytes = generate_acceptance_pdf(user_name, reference_id, package_type, current_date)
        pdf_base64 = base64.b64encode(pdf_bytes).decode('utf-8')
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{
                    font-family: 'Georgia', 'Times New Roman', serif;
                    background-color: #ffffff;
                    color: #333333;
                    padding: 40px;
                    margin: 0;
                    line-height: 1.8;
                }}
                .container {{
                    max-width: 650px;
                    margin: 0 auto;
                    background-color: #ffffff;
                    padding: 40px;
                }}
                .header {{
                    text-align: center;
                    margin-bottom: 40px;
                    padding-bottom: 20px;
                    border-bottom: 2px solid #a553be;
                }}
                .header h1 {{
                    color: #a553be;
                    font-size: 24px;
                    margin: 0;
                    font-weight: 600;
                }}
                .info-table {{
                    width: 100%;
                    margin-bottom: 30px;
                }}
                .info-table td {{
                    padding: 8px 0;
                    vertical-align: top;
                }}
                .info-label {{
                    color: #666666;
                    width: 180px;
                    font-weight: 500;
                }}
                .info-value {{
                    color: #333333;
                }}
                .content {{
                    margin-bottom: 30px;
                }}
                .content p {{
                    margin: 0 0 20px 0;
                }}
                .signature {{
                    margin-top: 40px;
                    padding-top: 20px;
                }}
                .signature p {{
                    margin: 2px 0;
                }}
                .signature .name {{
                    font-weight: 600;
                    color: #333333;
                }}
                .signature .title {{
                    color: #666666;
                    font-style: italic;
                }}
                .signature .company {{
                    color: #a553be;
                    font-weight: 500;
                }}
                .footer {{
                    margin-top: 40px;
                    padding-top: 20px;
                    border-top: 1px solid #e0e0e0;
                    text-align: center;
                    font-size: 12px;
                    color: #999999;
                }}
                .pdf-notice {{
                    background-color: #f8f8f8;
                    border: 1px solid #e0e0e0;
                    border-radius: 8px;
                    padding: 15px;
                    margin-top: 20px;
                    text-align: center;
                    font-size: 13px;
                    color: #666666;
                }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="header">
                    <h1>Acceptance of Submission Form</h1>
                </div>
                
                <table class="info-table">
                    <tr>
                        <td class="info-label">Date</td>
                        <td class="info-value">{current_date}</td>
                    </tr>
                    <tr>
                        <td class="info-label">Reference Number</td>
                        <td class="info-value">{reference_id[:8].upper()}</td>
                    </tr>
                </table>
                
                <div class="content">
                    <p>Dear {user_name},</p>
                    
                    <p>Your Profile submission has been well received by 2good2breal. Our team have begun the initial stages of verification and A.I. systems analysis.</p>
                    
                    <p>Once your <strong>{package_display}</strong> report has been thoroughly and conclusively completed by us, we will email it to you. Typically, this will be within 48 hours.</p>
                    
                    <p>For any queries you may have in the meantime, please email or call us directly during business hours.</p>
                </div>
                
                <div class="signature">
                    <p>Best regards,</p>
                    <br>
                    <p class="name">Jamie Madison</p>
                    <p class="title">Associate CEO</p>
                    <p class="company">2good2breal</p>
                </div>
                
                <div class="pdf-notice">
                    📎 A PDF copy of this acceptance form is attached to this email for your records.
                </div>
                
                <div class="footer">
                    <p>WhatsApp: +33 (0) 7 43 66 05 55 | Office: +33 (0) 7 67 92 55 45</p>
                </div>
            </div>
        </body>
        </html>
        """
        
        params = {
            "from": "2good2breal <onboarding@resend.dev>",
            "to": [user_email],
            "subject": f"Acceptance of Submission - Reference #{reference_id[:8].upper()}",
            "html": html_content,
            "attachments": [
                {
                    "filename": f"Acceptance_Form_{reference_id[:8].upper()}.pdf",
                    "content": pdf_base64
                }
            ]
        }
        
        await asyncio.to_thread(resend.Emails.send, params)
        logging.info(f"Client acceptance confirmation with PDF sent to: {user_email}")
    except Exception as e:
        logging.error(f"Failed to send client acceptance confirmation to {user_email}: {str(e)}")
        import traceback
        logging.error(f"Traceback: {traceback.format_exc()}")

@api_router.post("/analyze", response_model=VerificationResult)
async def analyze_profile(profile: ProfileAnalysisRequest, current_user: dict = Depends(get_current_user)):
    """Analyze a dating profile using AI to detect potential red flags."""
    
    # Check if user has PAID credits (free credits not accepted)
    basic_credits = current_user.get("basic_credits", 0)
    comprehensive_credits = current_user.get("comprehensive_credits", 0)
    premium_credits = current_user.get("premium_credits", 0)
    
    # Only count paid credits
    paid_credits = basic_credits + comprehensive_credits + premium_credits
    
    if paid_credits <= 0:
        raise HTTPException(
            status_code=402, 
            detail="No paid credits available. Please purchase a verification package to continue. Free credits are not accepted for profile analysis."
        )
    
    # Determine which credit to use (priority: basic > comprehensive > premium)
    credit_field = None
    credit_type = None
    if basic_credits > 0:
        credit_field = "basic_credits"
        credit_type = "basic"
    elif comprehensive_credits > 0:
        credit_field = "comprehensive_credits"
        credit_type = "comprehensive"
    elif premium_credits > 0:
        credit_field = "premium_credits"
        credit_type = "premium"
    
    # Count photos
    photos_count = len(profile.photos) if profile.photos else 0
    
    # Send form data by email to admin
    await send_analysis_form_notification(
        current_user["email"], 
        current_user["name"], 
        profile,
        photos_count
    )
    
    result_id = str(uuid.uuid4())
    
    # Run AI analysis for admin (in background, don't show to user)
    ai_analysis = None
    try:
        ai_analysis = await analyze_profile_with_ai_v2(profile)
        logging.info(f"AI analysis completed for profile: {profile.profile_name}")
    except Exception as e:
        logging.error(f"AI analysis failed: {e}")
        ai_analysis = {"error": str(e), "status": "failed"}
    
    # Create a confirmation result (no AI analysis shown to user)
    result = VerificationResult(
        id=result_id,
        user_id=current_user["id"],
        profile_name=profile.profile_name,
        overall_score=0,
        trust_level="pending",
        red_flags=[],
        analysis_summary="Your profile submission has been received successfully. Our team will analyze the profile and contact you within 48 hours with the results.",
        detailed_analysis={},
        image_analysis={},
        recommendations=["Your request is being processed by our expert team.", "You will receive the detailed verification report via email.", "For urgent inquiries, please contact us via WhatsApp or phone."],
        created_at=datetime.now(timezone.utc).isoformat()
    )
    
    # Save to database with AI analysis for admin
    result_dict = result.model_dump()
    result_dict["red_flags"] = []
    result_dict["credit_type_used"] = credit_type
    result_dict["status"] = "pending"
    result_dict["ai_analysis"] = ai_analysis  # Store AI analysis for admin
    result_dict["form_data"] = {
        "client_email": profile.client_email,
        "client_age": profile.client_age,
        "client_location": profile.client_location,
        "profile_name": profile.profile_name,
        "full_real_name": profile.full_real_name,
        "gender": profile.gender,
        "height": profile.height,
        "nationality": profile.nationality,
        "language_of_communication": profile.language_of_communication,
        "date_of_birth": profile.date_of_birth,
        "assumed_age": profile.assumed_age,
        "profile_location": profile.profile_location,
        "occupation": profile.occupation,
        "company_name": profile.company_name,
        "company_website": profile.company_website,
        "dating_platform": profile.dating_platform,
        "profile_bio": profile.profile_bio,
        "profile_photos_count": profile.profile_photos_count,
        "has_verified_photos": profile.has_verified_photos,
        "social_media_links": profile.social_media_links,
        "profile_creation_date": profile.profile_creation_date,
        "last_active": profile.last_active,
        "communication_frequency": profile.communication_frequency,
        "message_substance": profile.message_substance,
        "observations_concerns": profile.observations_concerns,
        "photos_uploaded": photos_count,
        "photos": [{"name": p.name, "base64": p.base64} for p in profile.photos] if profile.photos else []
    }
    await db.verification_results.insert_one(result_dict)
    
    # Deduct credit after successful submission
    await db.users.update_one(
        {"id": current_user["id"]},
        {
            "$inc": {credit_field: -1},
            "$set": {"updated_at": datetime.now(timezone.utc).isoformat()}
        }
    )
    
    # Send acceptance confirmation email to client
    # Use client_email if provided, otherwise fallback to user account email
    client_email = profile.client_email.strip() if profile.client_email else current_user["email"]
    if not client_email:
        client_email = current_user["email"]
    
    await send_client_acceptance_confirmation(
        client_email,
        current_user["name"],
        result_id,
        credit_type
    )
    
    logging.info(f"Analysis request submitted for user {current_user['email']}, acceptance sent to {client_email}, used 1 {credit_type} credit")
    
    return result

@api_router.get("/analyses", response_model=List[VerificationResult])
async def get_analyses(current_user: dict = Depends(get_current_user)):
    """Get all verification results for the current user."""
    results = await db.verification_results.find(
        {"user_id": current_user["id"]},
        {"_id": 0, "id": 1, "user_id": 1, "profile_name": 1, "overall_score": 1, "trust_level": 1, 
         "red_flags": 1, "analysis_summary": 1, "recommendations": 1, "created_at": 1, "status": 1}
    ).sort("created_at", -1).to_list(50)
    
    return results

@api_router.get("/analyses/{analysis_id}", response_model=VerificationResult)
async def get_analysis(analysis_id: str, current_user: dict = Depends(get_current_user)):
    """Get a specific verification result."""
    result = await db.verification_results.find_one(
        {"id": analysis_id, "user_id": current_user["id"]},
        {"_id": 0}
    )
    if not result:
        raise HTTPException(status_code=404, detail="Analysis not found")
    return result

@api_router.delete("/analyses/{analysis_id}")
async def delete_analysis(analysis_id: str, current_user: dict = Depends(get_current_user)):
    """Delete a verification result."""
    result = await db.verification_results.delete_one(
        {"id": analysis_id, "user_id": current_user["id"]}
    )
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Analysis not found")
    return {"message": "Analysis deleted"}

# ============== FILTER ROUTES ==============

@api_router.post("/filters", response_model=FilterResponse)
async def create_filter(filter_data: FilterCreate, current_user: dict = Depends(get_current_user)):
    """Create a custom verification filter."""
    filter_id = str(uuid.uuid4())
    
    filter_doc = {
        "id": filter_id,
        "user_id": current_user["id"],
        "name": filter_data.name,
        "description": filter_data.description or "",
        "criteria": filter_data.criteria,
        "is_active": filter_data.is_active,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.filters.insert_one(filter_doc)
    
    return FilterResponse(**filter_doc)

@api_router.get("/filters", response_model=List[FilterResponse])
async def get_filters(current_user: dict = Depends(get_current_user)):
    """Get all filters for the current user."""
    filters = await db.filters.find(
        {"user_id": current_user["id"]},
        {"_id": 0}
    ).to_list(100)
    
    return filters

@api_router.put("/filters/{filter_id}", response_model=FilterResponse)
async def update_filter(filter_id: str, filter_data: FilterCreate, current_user: dict = Depends(get_current_user)):
    """Update a custom filter."""
    existing = await db.filters.find_one(
        {"id": filter_id, "user_id": current_user["id"]},
        {"_id": 0}
    )
    if not existing:
        raise HTTPException(status_code=404, detail="Filter not found")
    
    update_data = {
        "name": filter_data.name,
        "description": filter_data.description or "",
        "criteria": filter_data.criteria,
        "is_active": filter_data.is_active
    }
    
    await db.filters.update_one(
        {"id": filter_id},
        {"$set": update_data}
    )
    
    updated = await db.filters.find_one({"id": filter_id}, {"_id": 0})
    return FilterResponse(**updated)

@api_router.delete("/filters/{filter_id}")
async def delete_filter(filter_id: str, current_user: dict = Depends(get_current_user)):
    """Delete a custom filter."""
    result = await db.filters.delete_one(
        {"id": filter_id, "user_id": current_user["id"]}
    )
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Filter not found")
    return {"message": "Filter deleted"}

# ============== DASHBOARD STATS ==============

@api_router.get("/stats")
async def get_stats(current_user: dict = Depends(get_current_user)):
    """Get dashboard statistics for the current user."""
    total_analyses = await db.verification_results.count_documents({"user_id": current_user["id"]})
    
    # Get trust level distribution
    pipeline = [
        {"$match": {"user_id": current_user["id"]}},
        {"$group": {"_id": "$trust_level", "count": {"$sum": 1}}}
    ]
    trust_distribution = await db.verification_results.aggregate(pipeline).to_list(10)
    
    # Get average score
    avg_pipeline = [
        {"$match": {"user_id": current_user["id"]}},
        {"$group": {"_id": None, "avg_score": {"$avg": "$overall_score"}}}
    ]
    avg_result = await db.verification_results.aggregate(avg_pipeline).to_list(1)
    avg_score = avg_result[0]["avg_score"] if avg_result else 0
    
    # Get recent analyses
    recent = await db.verification_results.find(
        {"user_id": current_user["id"]},
        {"_id": 0}
    ).sort("created_at", -1).limit(5).to_list(5)
    
    # Count total filters
    total_filters = await db.filters.count_documents({"user_id": current_user["id"]})
    
    return {
        "total_analyses": total_analyses,
        "average_score": round(avg_score, 1) if avg_score else 0,
        "trust_distribution": {item["_id"]: item["count"] for item in trust_distribution},
        "recent_analyses": recent,
        "total_filters": total_filters
    }

# ============== PAYMENT ROUTES ==============

@api_router.get("/packages", response_model=List[PackageInfo])
async def get_packages():
    """Get all available pricing packages."""
    return [
        PackageInfo(
            id=pkg_id,
            name=pkg["name"],
            amount=pkg["amount"],
            currency=pkg["currency"],
            description=pkg["description"],
            profiles_included=pkg["profiles_included"]
        )
        for pkg_id, pkg in PRICING_PACKAGES.items()
    ]

@api_router.post("/checkout", response_model=CheckoutResponse)
async def create_checkout(
    request: Request,
    checkout_data: CreateCheckoutRequest,
    current_user: dict = Depends(get_current_user)
):
    """Create a Stripe checkout session for a package."""
    import stripe
    
    # Validate package exists
    if checkout_data.package_id not in PRICING_PACKAGES:
        raise HTTPException(status_code=400, detail="Invalid package ID")
    
    package = PRICING_PACKAGES[checkout_data.package_id]
    
    # Build URLs from frontend origin
    success_url = f"{checkout_data.origin_url}/payment/success?session_id={{CHECKOUT_SESSION_ID}}"
    cancel_url = f"{checkout_data.origin_url}/payment/cancel"
    
    # Check for promo code and apply discount
    VALID_PROMO_CODE = "2good2026"
    DISCOUNT_PERCENT = 10
    
    original_amount = package["amount"]
    final_amount = original_amount
    promo_applied = False
    
    if checkout_data.promo_code and checkout_data.promo_code.lower() == VALID_PROMO_CODE.lower():
        final_amount = original_amount * (100 - DISCOUNT_PERCENT) / 100
        promo_applied = True
        logging.info(f"Promo code applied: {DISCOUNT_PERCENT}% discount for user {current_user['email']}")
    
    # Use native Stripe API with English locale
    stripe.api_key = STRIPE_API_KEY
    
    # Build product description with discount info if applicable
    product_description = package["description"]
    if promo_applied:
        product_description = f"{package['description']} (Code promo -{DISCOUNT_PERCENT}% appliqué)"
    
    session = stripe.checkout.Session.create(
        payment_method_types=["card"],
        line_items=[{
            "price_data": {
                "currency": package["currency"],
                "product_data": {
                    "name": package["name"],
                    "description": product_description
                },
                "unit_amount": int(final_amount * 100)  # Stripe uses cents
            },
            "quantity": 1
        }],
        mode="payment",
        success_url=success_url,
        cancel_url=cancel_url,
        locale="en",  # Force English language
        customer_email=current_user["email"],  # Pre-fill client's email on Stripe page
        metadata={
            "user_id": current_user["id"],
            "user_email": current_user["email"],
            "package_id": checkout_data.package_id,
            "package_name": package["name"],
            "profiles_included": str(package["profiles_included"]),
            "promo_code": checkout_data.promo_code if promo_applied else "",
            "original_amount": str(original_amount),
            "discount_percent": str(DISCOUNT_PERCENT) if promo_applied else "0"
        }
    )
    
    # Create payment transaction record in database
    transaction = {
        "id": str(uuid.uuid4()),
        "session_id": session.id,
        "user_id": current_user["id"],
        "user_email": current_user["email"],
        "package_id": checkout_data.package_id,
        "package_name": package["name"],
        "amount": final_amount,
        "original_amount": original_amount,
        "currency": package["currency"],
        "profiles_included": package["profiles_included"],
        "promo_code": checkout_data.promo_code if promo_applied else None,
        "discount_percent": DISCOUNT_PERCENT if promo_applied else 0,
        "payment_status": "pending",
        "status": "initiated",
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.payment_transactions.insert_one(transaction)
    
    logging.info(f"Checkout session created: {session.id} for user {current_user['email']}")
    
    return CheckoutResponse(
        checkout_url=session.url,
        session_id=session.id
    )

@api_router.get("/checkout/status/{session_id}", response_model=PaymentStatusResponse)
async def get_checkout_status(
    request: Request,
    session_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Get the status of a checkout session and update transaction."""
    
    # Find the transaction
    transaction = await db.payment_transactions.find_one(
        {"session_id": session_id, "user_id": current_user["id"]},
        {"_id": 0}
    )
    
    if not transaction:
        raise HTTPException(status_code=404, detail="Transaction not found")
    
    # Initialize Stripe checkout
    host_url = str(request.base_url).rstrip('/')
    webhook_url = f"{host_url}/api/webhook/stripe"
    stripe_checkout = StripeCheckout(api_key=STRIPE_API_KEY, webhook_url=webhook_url)
    
    # Get status from Stripe
    checkout_status: CheckoutStatusResponse = await stripe_checkout.get_checkout_status(session_id)
    
    # Update transaction if status changed and not already processed
    if transaction["payment_status"] != checkout_status.payment_status:
        update_data = {
            "payment_status": checkout_status.payment_status,
            "status": checkout_status.status,
            "updated_at": datetime.now(timezone.utc).isoformat()
        }
        
        # If payment is successful and not already credited
        if checkout_status.payment_status == "paid" and transaction["payment_status"] != "paid":
            # Add credits to user
            profiles_to_add = transaction["profiles_included"]
            credit_field = f"{transaction['package_id']}_credits"
            
            await db.users.update_one(
                {"id": current_user["id"]},
                {
                    "$inc": {credit_field: profiles_to_add},
                    "$set": {"updated_at": datetime.now(timezone.utc).isoformat()}
                }
            )
            
            update_data["credits_added"] = profiles_to_add
            logging.info(f"Added {profiles_to_add} {transaction['package_id']} credits to user {current_user['email']}")
        
        await db.payment_transactions.update_one(
            {"session_id": session_id},
            {"$set": update_data}
        )
    
    return PaymentStatusResponse(
        status=checkout_status.status,
        payment_status=checkout_status.payment_status,
        amount=transaction["amount"],
        currency=transaction["currency"],
        package_id=transaction["package_id"],
        package_name=transaction["package_name"]
    )

@api_router.get("/credits", response_model=UserCreditsResponse)
async def get_user_credits(current_user: dict = Depends(get_current_user)):
    """Get user's available credits."""
    user = await db.users.find_one({"id": current_user["id"]}, {"_id": 0})
    
    basic = user.get("basic_credits", 0)
    comprehensive = user.get("comprehensive_credits", 0)
    premium = user.get("premium_credits", 0)
    
    return UserCreditsResponse(
        basic_credits=basic,
        comprehensive_credits=comprehensive,
        premium_credits=premium,
        total_analyses_available=basic + comprehensive + (premium * 2)
    )

@api_router.get("/transactions")
async def get_user_transactions(current_user: dict = Depends(get_current_user)):
    """Get user's payment transaction history."""
    transactions = await db.payment_transactions.find(
        {"user_id": current_user["id"]},
        {"_id": 0}
    ).sort("created_at", -1).to_list(50)
    
    return transactions

@api_router.post("/webhook/stripe")
async def stripe_webhook(request: Request):
    """Handle Stripe webhooks."""
    try:
        body = await request.body()
        signature = request.headers.get("Stripe-Signature")
        
        host_url = str(request.base_url).rstrip('/')
        webhook_url = f"{host_url}/api/webhook/stripe"
        stripe_checkout = StripeCheckout(api_key=STRIPE_API_KEY, webhook_url=webhook_url)
        
        webhook_response = await stripe_checkout.handle_webhook(body, signature)
        
        if webhook_response.payment_status == "paid":
            # Update transaction
            session_id = webhook_response.session_id
            transaction = await db.payment_transactions.find_one(
                {"session_id": session_id},
                {"_id": 0}
            )
            
            if transaction and transaction["payment_status"] != "paid":
                # Add credits to user
                user_id = transaction["user_id"]
                profiles_to_add = transaction["profiles_included"]
                credit_field = f"{transaction['package_id']}_credits"
                
                await db.users.update_one(
                    {"id": user_id},
                    {
                        "$inc": {credit_field: profiles_to_add},
                        "$set": {"updated_at": datetime.now(timezone.utc).isoformat()}
                    }
                )
                
                await db.payment_transactions.update_one(
                    {"session_id": session_id},
                    {"$set": {
                        "payment_status": "paid",
                        "status": "complete",
                        "credits_added": profiles_to_add,
                        "updated_at": datetime.now(timezone.utc).isoformat()
                    }}
                )
                
                logging.info(f"Webhook: Added {profiles_to_add} credits to user {user_id}")
                
                # Send payment confirmation email to the CLIENT
                user = await db.users.find_one({"id": user_id}, {"_id": 0})
                if user:
                    await send_payment_confirmation_to_client(
                        user_email=transaction["user_email"],
                        user_name=user.get("name", "Customer"),
                        package_name=transaction["package_name"],
                        amount=transaction["amount"],
                        currency=transaction["currency"],
                        credits_added=profiles_to_add
                    )
        
        return {"status": "received"}
    except Exception as e:
        logging.error(f"Webhook error: {e}")
        raise HTTPException(status_code=400, detail=str(e))

# ============== HEALTH CHECK ==============

@api_router.get("/")
async def root():
    return {"message": "Profile Verify API", "status": "healthy"}

@api_router.get("/health")
async def health():
    """Health check endpoint for Kubernetes liveness/readiness probes."""
    health_status = {
        "status": "healthy",
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "service": "2good2breal-api"
    }
    
    # Check MongoDB connection if available
    if db is not None:
        try:
            await db.command('ping')
            health_status["database"] = "connected"
        except Exception as e:
            health_status["database"] = "disconnected"
            logger.warning(f"Database health check failed: {e}")
    else:
        health_status["database"] = "not_configured"
    
    return health_status

# Include the router in the main app
app.include_router(api_router)

# Root-level health endpoint (without /api prefix) for Kubernetes probes
@app.get("/health")
async def root_health():
    """Root health check endpoint for Kubernetes liveness/readiness probes."""
    return {"status": "healthy", "timestamp": datetime.now(timezone.utc).isoformat()}

@app.get("/")
async def app_root():
    """Root endpoint."""
    return {"message": "2good2breal API", "status": "running"}

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
